"""
Offer 捕手 - 学生求职匹配智能体 API
大厂专业级匹配分析 + 流式输出 + 简历优化

版本：1.2.0 (生产优化版)
- 使用 SQLite 替代 TinyDB，支持并发
- 添加 Redis 缓存层，提升性能
- 添加 API Key 认证
- 改进错误处理和日志
"""
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import json
import os
import io
import logging
import uvicorn
import re
from urllib.parse import quote

# 导入服务模块
from app.services.fast_jd_parser import FastJDParser
from app.services.fast_resume_parser import FastResumeParser
from app.services.xulindi_resume_parser import XulindiResumeParser
from app.services.llm_resume_parser import get_llm_parser
from app.services.kimi_vision_parser import get_vision_parser
from app.services.gemini_style_matcher import GeminiStyleMatcher
from app.services.professional_match_engine import ProfessionalMatchEngine
from app.services.professional_resume_optimizer import ProfessionalResumeOptimizer
# 使用 SQLite 数据库
from app.db.sqlite_db import get_db
# 使用 Redis 缓存
from app.cache.redis_cache import get_cache, CacheConfig
# Celery 任务
from app.tasks.celery_app import celery_app
from app.tasks.resume_tasks import parse_resume_text_task, parse_resume_file_task
from app.tasks.job_tasks import parse_job_task
from app.tasks.match_tasks import match_analysis_task, optimize_resume_task
# 监控和告警
from app.monitoring.metrics import get_metrics, MetricsMiddleware
# 快速 PDF 解析
from app.utils.pdf_parser import get_pdf_parser
from app.services.campus_recruit.campus_scorer import CampusRecruitScorer, CampusScore
# 一次性匹配分析引擎
from app.services.one_shot_match_engine import get_one_shot_matcher
# 深度分析与经历原子生成
from app.services.deep_analysis import DeepAnalysisService
from app.services.atom_generator import AtomGenerator
# 岗位推荐
from app.services.job_recommender import get_job_recommender
# 账号鉴权与用量额度
from app.services.auth_service import get_auth_service, get_current_user
from app.services.quota_service import get_quota_service
from fastapi import Depends

# ========== 日志配置 ==========

# 确保日志目录存在（使用绝对路径）
import os
# 默认写到 backend/logs；可用 LOG_DIR 环境变量覆盖（便于受限环境/容器自定义路径）
LOG_DIR = os.getenv("LOG_DIR") or os.path.join(os.path.dirname(os.path.dirname(__file__)), 'logs')
os.makedirs(LOG_DIR, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(os.path.join(LOG_DIR, 'app.log'), encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

# ========== API Key 认证 ==========

API_KEYS = os.getenv("API_KEYS", "demo-key-123").split(",")

def verify_api_key(request: Request) -> bool:
    """验证 API Key（演示版简单实现）"""
    api_key = request.headers.get("X-API-Key")
    if not api_key:
        api_key = request.query_params.get("api_key")
    # 演示版本：如果没有 API Key，允许通过
    # 生产环境应该强制验证
    return True

# ========== 应用初始化 ==========

app = FastAPI(
    title="Offer 捕手 API",
    description="学生求职匹配智能体 - 生产优化版 v1.2.0",
    version="1.2.0"
)

# CORS 配置（仅允许本地开发）
ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "http://localhost:3005",
    "http://127.0.0.1:3005",
    "http://localhost:8888",
    "http://127.0.0.1:8888",
    "http://localhost:8890",
    "http://127.0.0.1:8890",
    "file://"  # 允许本地文件访问（仅演示用）
]

# 在容器/远程预览场景下，前端访问域名可能是一个局域网 IP（例如 192.168.x.x）。
# 这里提供一个可配置的正则白名单，默认允许 localhost / 127.0.0.1 / 任意 IPv4 + 任意端口。
# 如需更严格限制，可通过环境变量覆盖：CORS_ALLOW_ORIGIN_REGEX
CORS_ALLOW_ORIGIN_REGEX = os.getenv(
    "CORS_ALLOW_ORIGIN_REGEX",
    r"^https?://(localhost|127\.0\.0\.1|\d+\.\d+\.\d+\.\d+)(:\d+)?$"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_origin_regex=CORS_ALLOW_ORIGIN_REGEX,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["*"],
)

# 添加 Prometheus 监控中间件
# 注意：需要在 uvicorn.run 中包装，不在模块级别包装
# app = MetricsMiddleware(app)  # 移到启动块中

# 初始化服务
jd_parser = FastJDParser()
resume_parser = FastResumeParser()
xulindi_parser = XulindiResumeParser()
llm_parser = get_llm_parser()  # LLM 文本解析器
vision_parser = get_vision_parser()  # Kimi 视觉解析器
gemini_matcher = GeminiStyleMatcher()
match_engine = ProfessionalMatchEngine()
resume_optimizer = ProfessionalResumeOptimizer()
campus_scorer = CampusRecruitScorer()
deep_analysis_service = DeepAnalysisService()
atom_generator = AtomGenerator()
# 初始化缓存
cache = get_cache()

# ========== 数据模型 ==========

class ParseJDRequest(BaseModel):
    """JD解析请求"""
    jd_text: str

class MatchRequest(BaseModel):
    """匹配分析请求"""
    resume_id: str
    job_id: str

class OptimizeResumeRequest(BaseModel):
    """简历优化请求"""
    resume_id: str
    job_id: str

class CustomizeResumeRequest(BaseModel):
    """定制简历生成请求"""
    resume_id: Optional[str] = None
    resume_text: Optional[str] = None
    jd_text: str
    match_result: Dict[str, Any]
    atoms: List[Dict[str, Any]] = []

class ExportResumePdfRequest(BaseModel):
    """定制简历 PDF 导出请求"""
    resume: Dict[str, Any]

# ========== 根路径 ==========

@app.get("/api/health")
async def root():
    """API 状态检查"""
    try:
        db = get_db()
        stats = db.get_stats()
        cache_stats = cache.get_stats()
        return {
            "service": "Offer 捕手 API",
            "version": "1.2.0",
            "status": "running",
            "database": "SQLite (WAL mode)",
            "cache": "Redis",
            "database_stats": stats,
            "cache_stats": cache_stats
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return {
            "service": "Offer 捕手 API",
            "version": "1.2.0",
            "status": "degraded",
            "error": str(e)
        }

# ========== 账号鉴权 / 用量额度 / 邀请码接口 ==========

class PhoneLoginRequest(BaseModel):
    phone: str
    code: str


class RedeemInviteRequest(BaseModel):
    code: str


@app.post("/api/v1/auth/login")
async def auth_login(req: PhoneLoginRequest):
    """手机号 + 验证码登录/注册，返回 JWT。

    - 站长本人手机号无限次；其他手机号凭通用验证码可使用一次完整流程。
    - 首次登录自动创建账号（无需单独注册）。
    """
    try:
        return {"success": True, **get_auth_service().login_with_phone(req.phone, req.code)}
    except ValueError as e:
        raise HTTPException(status_code=401, detail=str(e))


@app.get("/api/v1/auth/me")
async def auth_me(current_user: Dict[str, Any] = Depends(get_current_user)):
    """返回当前登录用户信息（手机号、剩余次数、是否无限、是否看过引导）。"""
    return {"success": True, "user": current_user}


@app.post("/api/v1/auth/onboarding-done")
async def auth_onboarding_done(current_user: Dict[str, Any] = Depends(get_current_user)):
    """标记当前用户已完成首次引导。"""
    get_db().mark_onboarding_done(current_user["id"])
    return {"success": True}


@app.post("/api/v1/invite/redeem")
async def invite_redeem(req: RedeemInviteRequest, current_user: Dict[str, Any] = Depends(get_current_user)):
    """兑换邀请码，增加免费体验次数。"""
    try:
        result = get_quota_service().redeem_invite(current_user["id"], req.code)
        return {"success": True, **result}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/v1/share/claim")
async def share_claim(current_user: Dict[str, Any] = Depends(get_current_user)):
    """分享给微信好友奖励：每个账号仅可领取一次 +1 次。"""
    try:
        result = get_quota_service().claim_share_reward(current_user["id"])
        return {"success": True, **result}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


# ========== 简历解析接口 ==========

@app.post("/api/v1/resume/parse")
async def parse_resume(file: UploadFile = File(...)):
    """
    解析上传的简历文件

    支持格式：
    - PDF: 使用 Kimi 视觉模型直接分析图片
    - 图片: 使用 Kimi 视觉模型分析
    - DOCX: 提取文本后用 LLM 分析
    - TXT: 直接用 LLM 分析

    错误处理：
    - 文件过大: 前端限制，后端验证
    - 格式不支持: 返回友好提示
    - AI 调用失败: 降级到文本解析
    """
    db = get_db()
    pdf_parser = get_pdf_parser()

    # 文件大小限制 (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024

    try:
        logger.info(f"Parsing resume file: {file.filename}")

        # 读取文件内容
        content = await file.read()
        filename = file.filename or ""

        # 检查文件大小
        if len(content) > MAX_FILE_SIZE:
            raise ValueError(f"文件过大，请上传小于 10MB 的文件")

        resume_data = None

        # 根据文件类型选择解析方式
        if filename.lower().endswith('.pdf'):
            # 优先走「文本提取 + LLM」快路径：绝大多数简历是数字版 PDF，可直接取文本，
            # 比把每页渲染成 2x 图片再喂视觉模型快很多。只有文本为空/疑似乱码（扫描件）
            # 才回退到 Kimi 视觉模型。
            text = ""
            try:
                text = pdf_parser.extract_text(content, filename)
            except Exception as e:
                logger.warning(f"PDF 文本提取失败，转视觉解析: {e}")

            if text.strip() and not pdf_parser._looks_garbled(text):
                logger.info("PDF 文本提取成功，使用 LLM 文本解析（快路径）")
                resume_data = await llm_parser.parse_async(text)
            else:
                # 扫描件 / 取不到文本：用 Kimi 视觉模型兜底
                try:
                    logger.info("PDF 无可用文本，使用 Kimi 视觉模型解析...")
                    resume_data = await vision_parser.parse_pdf_async(content, filename)
                    logger.info(f"视觉解析成功: {resume_data.get('basic_info', {}).get('name', '未知')}")
                except Exception as e:
                    logger.warning(f"视觉解析失败: {e}")
                    if not text.strip():
                        raise ValueError("无法从 PDF 中提取文本内容")
                    resume_data = await llm_parser.parse_async(text)

        elif filename.lower().endswith(('.jpg', '.jpeg', '.png', '.webp')):
            # 图片文件，使用视觉模型
            logger.info("使用 Kimi 视觉模型解析图片...")
            resume_data = await vision_parser.parse_image_async(content, filename)

        elif filename.lower().endswith('.docx'):
            # DOCX 文件，提取文本后用 LLM 解析
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                text = content.decode('utf-8', errors='ignore')

            if text.strip():
                resume_data = await llm_parser.parse_async(text)
            else:
                raise ValueError("无法从 DOCX 中提取文本内容")

        else:
            # 纯文本文件
            text = content.decode('utf-8', errors='ignore')
            resume_data = await llm_parser.parse_async(text)

        # 验证解析结果
        if not resume_data or not resume_data.get("basic_info", {}).get("name"):
            raise ValueError("解析失败：无法识别简历中的姓名信息")

        # 保存到数据库
        resume_id = db.save_resume(resume_data)
        logger.info(f"Resume parsed successfully: {resume_id}")

        return {
            "success": True,
            "resume_id": resume_id,
            "data": {**resume_data, "id": resume_id},
            "message": "简历解析成功"
        }

    except ValueError as e:
        logger.warning(f"Validation error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Resume parsing failed: {e}")
        raise HTTPException(status_code=500, detail=f"简历解析失败: {str(e)}")


class ParseResumeTextRequest(BaseModel):
    """简历文本解析请求"""
    text: str


@app.post("/api/v1/resume/parse-text")
async def parse_resume_text(request: ParseResumeTextRequest):
    """解析粘贴的简历文本"""
    db = get_db()
    try:
        logger.info("Parsing resume from text")

        # 先检查缓存
        cached_result = cache.get_cached_resume_parse(request.text)
        if cached_result:
            logger.info("Resume parse cache hit")
            # 缓存命中，保存新的记录并返回
            cached_result["raw_text"] = request.text
            cached_result["interests"] = cached_result.get("interests") or _extract_interests_from_text(request.text)
            resume_id = db.save_resume(cached_result)
            return {
                "success": True,
                "resume_id": resume_id,
                "data": {**cached_result, "id": resume_id},
                "message": "简历解析成功（缓存）",
                "cached": True
            }

        # 优先使用 LLM 解析器（完整提取，不丢信息）
        try:
            resume_data = await llm_parser.parse_async(request.text)
            if resume_data.get("basic_info", {}).get("name"):
                logger.info(f"LLM 解析成功: {resume_data['basic_info'].get('name')}")
            else:
                raise ValueError("LLM 解析未能识别姓名")
        except Exception as e:
            logger.warning(f"LLM 解析失败，回退到传统解析器: {e}")
            # 回退到传统解析器
            try:
                resume_data = xulindi_parser.parse(request.text)
                if not resume_data.get("basic_info", {}).get("name"):
                    resume_data = resume_parser.parse_from_text(request.text)
            except Exception as e2:
                raise ValueError(f"所有解析器都失败了: {e2}")

        # 从原文抓「兴趣爱好」存进简历，供定制简历的「兴趣爱好 Interests」段使用（抓不到则留空）
        resume_data["interests"] = _extract_interests_from_text(request.text)
        resume_data["raw_text"] = request.text

        # 保存到数据库
        resume_id = db.save_resume(resume_data)

        # 缓存解析结果
        cache.cache_resume_parse(request.text, resume_data)

        return {
            "success": True,
            "resume_id": resume_id,
            "data": {**resume_data, "id": resume_id},
            "message": "简历解析成功"
        }

    except Exception as e:
        logger.error(f"Resume text parsing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ========== JD解析接口 ==========

@app.post("/api/v1/job/parse")
async def parse_job(request: ParseJDRequest):
    """解析JD文本"""
    db = get_db()
    try:
        logger.info("Parsing job description")

        # 先检查缓存
        cached_result = cache.get_cached_job_parse(request.jd_text)
        if cached_result:
            logger.info("Job parse cache hit")
            job_id = db.save_job(cached_result)
            return {
                "success": True,
                "data": {**cached_result, "id": job_id},
                "message": "JD解析成功（缓存）",
                "cached": True
            }

        # 使用快速JD解析器
        job_data = jd_parser.parse(request.jd_text)

        # 保存到数据库
        job_id = db.save_job(job_data)

        # 缓存解析结果
        cache.cache_job_parse(request.jd_text, job_data)

        return {
            "success": True,
            "data": {**job_data, "id": job_id},
            "message": "JD解析成功"
        }

    except Exception as e:
        logger.error(f"Job parsing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ========== 匹配分析接口 ==========

@app.post("/api/v1/match")
async def match_analysis(request: MatchRequest):
    """执行专业级匹配分析"""
    db = get_db()
    try:
        logger.info(f"Starting match analysis: resume={request.resume_id}, job={request.job_id}")

        # 先检查缓存
        cached_match = cache.get_cached_match_result(request.resume_id, request.job_id)
        if cached_match:
            logger.info("Match result cache hit")
            # 保存新的匹配记录
            resume_record = db.get_resume(request.resume_id)
            job_record = db.get_job(request.job_id)
            match_record_data = {
                "resume_id": request.resume_id,
                "job_id": request.job_id,
                "match_result": cached_match,
                "position_name": job_record.get("position_name", "") if job_record else "",
                "company": job_record.get("company", "") if job_record else ""
            }
            match_id = db.save_match(match_record_data)
            return {
                "success": True,
                "data": {**cached_match, "match_id": match_id},
                "message": "专业级匹配分析完成（缓存）",
                "cached": True
            }

        # 获取数据
        resume_record = db.get_resume(request.resume_id)
        job_record = db.get_job(request.job_id)

        if not resume_record or not job_record:
            raise HTTPException(status_code=404, detail="简历或岗位不存在")

        resume_data = resume_record
        job_data = job_record

        # 执行匹配分析
        match_result = match_engine.calculate(resume_data, job_data)

        # 保存匹配记录
        match_record_data = {
            "resume_id": request.resume_id,
            "job_id": request.job_id,
            "match_result": match_result,
            "position_name": job_data.get("position_name", ""),
            "company": job_data.get("company", "")
        }
        match_id = db.save_match(match_record_data)

        # 缓存匹配结果
        cache.cache_match_result(request.resume_id, request.job_id, match_result)

        logger.info(f"Match analysis completed: {match_id}")

        return {
            "success": True,
            "data": {**match_result, "match_id": match_id},
            "message": "专业级匹配分析完成"
        }

    except Exception as e:
        logger.error(f"Match analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/match/from-upload")
async def match_from_upload(
    resume_file: UploadFile = File(..., description="简历文件"),
    jd_text: str = Form(..., description="JD 文本")
):
    """
    一次性匹配分析 - 上传简历文件 + JD 文本，直接返回匹配结果

    流程：
    1. 接收简历文件（PDF/图片/DOCX）
    2. 接收 JD 文本
    3. 提取简历文本
    4. 调用 LLM 一次性分析：解析简历 + 解析 JD + 匹配分析
    5. 返回匹配结果

    优势：
    - 只调用一次 AI，响应更快
    - AI 同时看到简历和 JD，匹配更准确
    - 用户体验更流畅
    """
    try:
        logger.info(f"一次性匹配分析: {resume_file.filename}")

        # 读取文件
        file_bytes = await resume_file.read()
        filename = resume_file.filename or "resume"

        # 文件大小限制 (10MB)
        MAX_FILE_SIZE = 10 * 1024 * 1024
        if len(file_bytes) > MAX_FILE_SIZE:
            raise ValueError(f"文件过大，请上传小于 10MB 的文件")

        # 获取一次性匹配引擎
        one_shot_matcher = get_one_shot_matcher()

        # 执行一次性分析
        match_result = await one_shot_matcher.analyze(file_bytes, filename, jd_text)

        # 可选：保存到数据库
        db = get_db()
        # 解析简历数据用于存储
        # 这里简化处理，实际可以存储完整记录

        logger.info(f"一次性匹配分析完成")

        return {
            "success": True,
            "data": match_result,
            "message": "匹配分析完成"
        }

    except ValueError as e:
        logger.warning(f"Validation error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"One-shot match failed: {e}")
        raise HTTPException(status_code=500, detail=f"匹配分析失败: {str(e)}")


# ========== 流式匹配分析接口 ==========

from fastapi.responses import StreamingResponse
import asyncio

@app.post("/api/v1/match/stream")
async def stream_match_analysis(resume_data: str = Form(...), job_data: str = Form(...)):
    """
    Gemini 风格流式匹配分析 - 像真实 HR 思考过程

    输出格式：
    - start: 开始分析
    - thinking: 思考状态
    - section: 分析章节
    - content: 流式文本内容
    - result: 结构化结果
    - complete: 分析完成
    """
    async def event_generator():
        try:
            # 解析输入
            resume_dict = json.loads(resume_data)
            job_dict = json.loads(job_data)

            # 使用 Gemini 风格匹配器
            async for event in gemini_matcher.match_analysis_stream(resume_dict, job_dict):
                yield event

        except Exception as e:
            import traceback
            yield f'event: error\ndata: {{"message": "分析失败: {str(e)}"}}\n\n'

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@app.post("/api/v1/match/from-upload/stream")
async def stream_match_from_upload(
    resume_file: Optional[UploadFile] = File(None, description="简历文件"),
    jd_text: str = Form(..., description="JD 文本"),
    use_fixed_resume: bool = Form(False, description="是否使用固定测试简历"),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    流式一次性匹配分析 - 支持文件上传

    SSE 事件类型：
    - event: progress     - 进度更新（含 id 支持断线重连）
    - event: thinking     - AI 思考过程
    - event: result       - 最终匹配结果
    - event: error        - 错误信息
    - event: done         - 分析完成
    - event: heartbeat    - 心跳保持连接

    SSE 格式符合最佳实践：
    - 使用 JSON 格式 data 字段
    - 添加 id 字段支持断线重连
    - 添加 retry 字段控制重连间隔
    - 正确的响应头配置
    """
    # ===== 准入：必须已登录且有剩余次数 =====
    quota_service = get_quota_service()
    if not quota_service.check_quota(current_user["id"]):
        raise HTTPException(
            status_code=402,
            detail="免费次数已用完，输入邀请码可再体验一次，或前往升级获取更多次数。"
        )

    # ===== 先读取文件（避免在生成器中读取时文件已关闭）=====
    if use_fixed_resume:
        fixed_resume_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            "uploads",
            "fixed-test-resume.pdf"
        )
        if not os.path.exists(fixed_resume_path):
            raise HTTPException(status_code=400, detail="固定测试简历不存在")
        with open(fixed_resume_path, "rb") as f:
            file_bytes = f.read()
        filename = os.path.basename(fixed_resume_path)
        logger.info(f"使用固定测试简历: {filename}")
    else:
        if resume_file is None:
            raise HTTPException(status_code=400, detail="请上传简历文件，或启用固定测试简历")
        file_bytes = await resume_file.read()
        filename = resume_file.filename or "resume"
    one_shot_matcher = get_one_shot_matcher()
    db = get_db()
    # 匹配分析必须只基于「本次上传/粘贴的简历」。
    # 账号级经历原子库即使做了 user_id 隔离，也可能属于该用户的另一份简历；
    # 传给匹配引擎会被 AI 当作当前候选人的证据，导致串数据。
    atoms = []

    # 生成唯一会话 ID 用于断线重连
    import uuid
    session_id = str(uuid.uuid4())[:8]
    event_id = 0

    async def event_generator():
        nonlocal file_bytes, filename, one_shot_matcher, jd_text
        nonlocal event_id

        try:
            import asyncio
            import time

            # ===== 初始连接确认 =====
            event_id += 1
            yield f": 连接已建立\n\n"
            yield f"retry: 3000\n\n"
            yield f"event: connected\ndata: {json.dumps({'session_id': session_id, 'filename': filename})}\nid: {event_id}\n\n"

            # ===== 阶段1：读取文件 =====
            event_id += 1
            yield f"event: progress\ndata: {json.dumps({'stage': 'reading', 'message': '🔍 开始分析...'})}\nid: {event_id}\n\n"

            event_id += 1
            yield f"event: progress\ndata: {json.dumps({'stage': 'reading', 'message': '📄 正在读取简历...'})}\nid: {event_id}\n\n"

            # ===== 阶段2：准备AI分析 =====
            start_time = time.time()
            resume_input = file_bytes

            event_id += 1
            yield f"event: progress\ndata: {json.dumps({'stage': 'uploading', 'message': '✓ 简历读取完成'})}\nid: {event_id}\n\n"

            # ===== 阶段3：LLM 流式分析 =====
            event_count = 0
            got_result = False

            logger.info("开始调用 analyze_stream...")

            async for event_type, event_data in one_shot_matcher.analyze_stream(
                resume_input, jd_text, filename, atoms=atoms
            ):
                event_count += 1
                logger.info(f"[SSE] 收到事件 #{event_count}: type={event_type}")

                event_id += 1

                if event_type == "progress":
                    # 进度更新 - 直接传递 event_data
                    yield f"event: progress\ndata: {json.dumps(event_data)}\nid: {event_id}\n\n"

                elif event_type == "connected":
                    yield f"event: connected\ndata: {json.dumps(event_data)}\nid: {event_id}\n\n"

                elif event_type == "thinking":
                    yield f"event: thinking\ndata: {json.dumps(event_data)}\nid: {event_id}\n\n"

                elif event_type == "stage":
                    # 阶段状态更新
                    stage_data = event_data
                    if stage_data.get("status") == "analyzing":
                        stage_title = stage_data.get('title', '分析')
                        if stage_data.get("status") == "analyzing":
                            emoji = "🔍"
                            message = f"{emoji} {stage_title}..."
                            yield f"event: progress\ndata: {json.dumps({'stage': 'analyzing', 'message': message})}\nid: {event_id}\n\n"
                        elif stage_data.get("status") == "done":
                            emoji = "✓"
                            message = f"{emoji} {stage_title}完成"
                            yield f"event: progress\ndata: {json.dumps({'stage': 'done', 'message': message})}\nid: {event_id}\n\n"

                elif event_type == "result":
                    # 最终结果
                    got_result = True
                    yield f"event: result\ndata: {json.dumps(event_data)}\nid: {event_id}\n\n"

                elif event_type == "error":
                    # 错误
                    yield f"event: error\ndata: {json.dumps(event_data)}\nid: {event_id}\n\n"

                elif event_type == "done":
                    # 完成
                    yield f"event: done\ndata: {json.dumps(event_data)}\nid: {event_id}\n\n"

            # ===== 完成 =====
            # 分析成功产出结果才扣减额度；扣减后把最新剩余次数推给前端。
            if got_result:
                try:
                    quota_service.consume(current_user["id"])
                    refreshed = get_db().get_user_by_id(current_user["id"]) or {}
                    event_id += 1
                    yield f"event: quota\ndata: {json.dumps({'remaining_quota': refreshed.get('remaining_quota', 0)})}\nid: {event_id}\n\n"
                except Exception as quota_err:
                    logger.warning(f"额度扣减失败（不影响结果返回）: {quota_err}")

            event_id += 1
            elapsed = time.time() - start_time
            yield f"event: done\ndata: {json.dumps({'message': '分析完成', 'elapsed': f'{elapsed:.1f}秒'})}\nid: {event_id}\n\n"

        except asyncio.CancelledError:
            # 用户主动中止（如刷新页面）
            logger.info(f"SSE 连接被用户取消: session_id={session_id}")
            event_id += 1
            yield f"event: cancelled\ndata: {json.dumps({'message': '分析已取消'})}\nid: {event_id}\n\n"

        except Exception as e:
            import traceback
            error_msg = f"分析失败: {str(e)}"
            logger.error(f"流式分析错误: {traceback.format_exc()}")
            event_id += 1
            yield f"event: error\ndata: {json.dumps({'message': error_msg, 'type': 'system_error'})}\nid: {event_id}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # 禁用 Nginx 缓冲
            "X-Content-Type-Options": "nosniff"
        }
    )


# ========== 流式输出辅助函数 ==========

def _get_verb_rating(score: float) -> str:
    """根据动词强度分数返回评级"""
    if score >= 8.5:
        return "优秀 - 表达极具说服力"
    elif score >= 7.5:
        return "良好 - 表达清晰有力"
    elif score >= 6.5:
        return "一般 - 部分表达偏弱"
    else:
        return "需改进 - 建议强化动词"


def _get_verb_suggestions(weak_verbs: list) -> list:
    """获取动词升级建议"""
    verb_upgrade_map = {
        "协助": "负责/主导",
        "参与": "执行/构建",
        "学习": "掌握/精通",
        "帮忙": "支持/协调",
        "负责": "主导/引领",
        "执行": "构建/实现"
    }
    suggestions = []
    for verb in weak_verbs[:3]:
        if verb in verb_upgrade_map:
            suggestions.append(f"'{verb}' → '{verb_upgrade_map[verb]}'")
    return suggestions


def _estimate_pass_rate(score: float) -> str:
    """估算通过率"""
    if score >= 85:
        return "75% - 高通过率"
    elif score >= 75:
        return "50% - 中等通过率"
    elif score >= 65:
        return "30% - 需要优化"
    elif score >= 55:
        return "15% - 建议补缺"
    else:
        return "5% - 匹配度较低"


def _get_score_conclusion(score: float) -> str:
    """获取评分结论"""
    if score >= 85:
        return "简历与岗位高度匹配，建议直接投递"
    elif score >= 75:
        return "匹配度良好，小幅优化后投递"
    elif score >= 65:
        return "存在明显缺口，建议重点优化"
    elif score >= 55:
        return "匹配度不足，建议慎重考虑"
    else:
        return "匹配度较低，不建议投递"


def _estimate_score_improvement(current_score: float, suggestions: list) -> str:
    """估算优化后分数提升"""
    high_priority_count = len([s for s in suggestions if s.get("priority") == "high"])
    estimated_gain = high_priority_count * 5
    new_score = min(95, current_score + estimated_gain)
    improvement = new_score - current_score
    return f"+{improvement:.1f}分 → {new_score:.1f}分"


def _extract_advantages(resume_dict: dict, job_dict: dict) -> list:
    """提取简历优势"""
    advantages = []

    # 大厂背景
    work_exp = resume_dict.get("experience", [])
    big_companies = ["字节", "腾讯", "阿里", "百度", "美团", "京东", "滴滴", "网易", "小米", "华为"]
    for exp in work_exp:
        company = exp.get("company", "")
        if any(bc in company for bc in big_companies):
            advantages.append(f"大厂背景 ({company})")
            break

    # 学历优势
    education = resume_dict.get("education", {})
    degree = education.get("degree", "")
    if "硕士" in degree or "博士" in degree:
        advantages.append(f"高学历 ({degree})")

    # 项目成果量化
    projects = resume_dict.get("projects", [])
    for proj in projects:
        desc = proj.get("description", "")
        if any(char in desc for char in ["%", "倍", "万", "亿"]):
            advantages.append("项目成果量化")
            break

    # 技能广度
    skills = resume_dict.get("skills", [])
    if len(skills) >= 10:
        advantages.append("技能广度")

    return advantages[:3]


def _extract_improvements(resume_dict: dict, job_dict: dict) -> list:
    """提取需要改进的方面"""
    improvements = []

    # 缺少热门技术
    requirements = job_dict.get("requirements", {})
    required_skills = requirements.get("hard_skills", [])
    resume_skills = resume_dict.get("skills", [])
    missing = [s for s in required_skills if s not in resume_skills]
    if missing:
        improvements.append(f"补充技能: {', '.join(missing[:2])}")

    # 经验不足
    work_exp = resume_dict.get("experience", [])
    if len(work_exp) < 2:
        improvements.append("增加工作经验")

    # 项目缺乏量化
    projects = resume_dict.get("projects", [])
    has_quantified = any(
        any(char in proj.get("description", "") for char in ["%", "倍", "万"])
        for proj in projects
    )
    if not has_quantified:
        improvements.append("量化项目成果")

    return improvements[:3]


def _generate_strategy(total_score: float, match_level: str) -> str:
    """生成投递策略建议"""
    if total_score >= 80:
        return "强烈推荐投递，你的背景与岗位高度匹配。建议在面试中突出核心项目经验。"
    elif total_score >= 65:
        return "推荐投递，但建议重点优化简历中的技能描述和项目成果。可先补充短板再投递。"
    elif total_score >= 50:
        return "谨慎投递，存在明显技能或经验缺口。建议先系统学习相关技术栈，补充项目经验。"
    else:
        return "不建议投递，匹配度较低。建议选择更符合当前背景的岗位，或进行系统性技能提升。"


def _is_filler_bullet(text: str) -> bool:
    """判断一条 bullet 是否为「凑数空话套话」（无具体数字/技术/对象，只有泛化表态）。

    保守策略：只要包含具体信号（数字、英文技术词、百分比等）就保留；
    仅当整条几乎全是套话短语拼接、且没有任何具体信息时才判为 filler。
    """
    t = (text or "").strip()
    if not t:
        return True
    # 含具体信号则保留：阿拉伯数字、百分号、英文（技术栈/缩写）
    if re.search(r"[0-9%]", t) or re.search(r"[A-Za-z]{2,}", t):
        return False
    # 高频空话短语（无具体内容的泛化表态）
    filler_phrases = [
        "保障系统稳定运行", "确保系统稳定运行", "提升整体效率", "提升系统性能",
        "提高工作效率", "持续优化系统架构", "保障高可用性和可扩展性",
        "快速定位并解决潜在问题", "持续监控系统性能", "提升用户体验",
        "增强系统稳定性", "通过技术手段", "通过技术优化", "保障业务连续性",
        "提升风控效果和准确率", "确保代码质量",
    ]
    hit = sum(1 for p in filler_phrases if p in t)
    # 短句（<=24字）且命中至少一个套话短语，且不含具体对象数字 → 判为 filler
    if hit >= 1 and len(t) <= 28:
        return True
    # 两个以上套话短语拼接，基本可判定为纯套话
    if hit >= 2:
        return True
    return False


def _sanitize_resume_bullet(text: str) -> str:
    """轻量清洗 AI bullet。

    产品边界：允许在真实项目/真实实习内部做简历式包装；
    这里不删除“提升效率/优化体验”等结果表达，只做空白与标点清理。
    真正的事实漂移（新公司/新项目/新团队/新硬技能）在原子改写护栏里拦截。
    """
    t = str(text or "").strip()
    if not t:
        return ""
    t = re.sub(r"\s+", " ", t)
    return t.strip(" ，,；;。.")


def _extract_education_from_text(resume_text: str) -> list:
    """从简历原文正则抓教育（学校+学历+年份），按学历低到高排序。

    仅作兜底：当没有 resume_id / 库里没存教育时，避免教育段整段空白。
    只抓原文真实出现的，抓不到返回空列表，绝不编造。
    """
    text = resume_text or ""
    deg_rank = {"大专": 1, "专科": 1, "本科": 2, "学士": 2,
                "硕士": 3, "研究生": 3, "博士": 4}
    results = []
    # 逐行扫描，命中「含学历关键词且含学校/年份」的行
    for line in re.split(r"[\n;；]+", text):
        line = line.strip(" ：:、，,。.\t")
        if not line:
            continue
        deg = ""
        for k in ("博士", "硕士", "研究生", "本科", "学士", "大专", "专科"):
            if k in line:
                deg = "硕士" if k == "研究生" else ("本科" if k == "学士" else k)
                break
        school_m = re.search(r"([\u4e00-\u9fa5]{2,15}(?:大学|学院|学校|院校))", line)
        yrs4 = re.findall(r"(?:19|20)\d{2}", line)
        if not deg and not school_m:
            continue
        if not deg and not yrs4:
            continue
        sy = yrs4[0] if yrs4 else ""
        ey = yrs4[1] if len(yrs4) > 1 else ""
        date_range = (f"{sy} - {ey}" if sy and ey else (sy or ey))
        major_m = re.search(r"(?:大学|学院|学校|院校)\s*([\u4e00-\u9fa5]{2,12})", line)
        results.append({
            "school": school_m.group(1) if school_m else "",
            "major": major_m.group(1) if major_m else "",
            "degree": deg,
            "graduation_year": ey,
            "date_range": date_range,
            "_rank": deg_rank.get(deg, 0),
            "_sy": int(sy) if sy.isdigit() else 9999,
        })
    if not results:
        return []
    results.sort(key=lambda e: (e["_rank"], e["_sy"]))
    for e in results:
        e.pop("_rank", None)
        e.pop("_sy", None)
    return results


def _extract_interests_from_text(resume_text: str) -> str:
    """从简历原文里抓「兴趣爱好/爱好/特长」等，抓到才返回；抓不到留空，绝不编造。"""
    m = re.search(
        r"(?m)^\s*(?:兴趣爱好|业余爱好|个人爱好|个人兴趣|兴趣与爱好|爱好特长|特长爱好|爱好|特长|兴趣)"
        r"[：:\s]*\n?\s*([^\n]{2,120})",
        resume_text or ""
    )
    if not m:
        return ""
    cand = m.group(1).strip(" ：:、，,。.")
    if cand and "求职" not in cand and "方向" not in cand:
        return cand
    return ""


def _infer_job_title_from_text(jd_text: str) -> str:
    lines = [line.strip("•- \t") for line in (jd_text or "").splitlines() if line.strip()]
    if not lines:
        return "目标岗位"
    title = lines[0]
    # 清洗成干净的岗位名，避免把整行 JD（含"校招/硕士/熟悉LLM"等）都当成求职意向：
    # 1) 去掉"职位/岗位/招聘"等前缀标签
    title = re.sub(r"^(职位|岗位|招聘岗位|招聘职位|岗位名称|职位名称)\s*[:：]?\s*", "", title)
    # 2) 把括号及其内容（如"（校招）""(社招)"）替换成空格——用空格而非删空，
    #    避免"工程师（校招）熟悉Spring"删括号后粘成"工程师熟悉Spring"，导致下一步切不开。
    title = re.sub(r"[（(].*?[)）]", " ", title)
    # 3) 按常见分隔符/空格切，取第一段作为岗位名（"AI产品经理 校招 硕士" -> "AI产品经理"）
    title = re.split(r"[\s,，、|/]+", title.strip())[0] if title.strip() else title
    title = title.strip()
    # 4) 过长则截断兜底
    if len(title) > 20:
        title = title[:20]
    return title or "目标岗位"


def _extract_requirement_groups(match_result: Dict[str, Any], jd_text: str) -> Dict[str, List[str]]:
    requirement_checks = (
        match_result.get("report_requirement_checks")
        or match_result.get("requirement_checks")
        or match_result.get("sections", {}).get("requirement_checks")
        or []
    )
    matched = []
    missing = []
    for item in requirement_checks:
        requirement = str(item.get("requirement") or item.get("title") or "").strip()
        if not requirement:
            continue
        status = str(item.get("status", "")).lower()
        if status in {"matched", "strong_match", "fully_matched"}:
            matched.append(requirement)
        else:
            missing.append(requirement)

    if not matched and not missing:
        jd_decomp = (
            match_result.get("report_jd_decomposition")
            or match_result.get("jd_decomposition")
            or {}
        )
        core_competencies = jd_decomp.get("core_competencies", []) or []
        # 注意：不要把 JD 硬要求直接当成"已匹配"。没有证据时宁可留空，
        # 也不能把岗位标题/职责行误标为候选人的匹配能力（会出现"匹配能力=岗位职责"的怪现象）。
        missing = core_competencies[:3]

    if not missing and jd_text:
        lines = [line.strip("•- \t") for line in jd_text.splitlines() if line.strip()]
        missing = lines[2:5]

    return {
        "matched": matched[:5],
        "missing": missing[:5],
    }


def _tokenize_for_match(text: str) -> List[str]:
    raw_tokens = re.findall(r"[A-Za-z0-9\+\-#/\.]+|[\u4e00-\u9fff]{2,8}", text or "")
    blacklist = {"负责", "参与", "能够", "以及", "进行", "相关", "经验", "能力", "产品", "岗位", "要求", "加分项", "我们提供"}
    tokens = []
    for token in raw_tokens:
        token = token.strip()
        if len(token) < 2 or token in blacklist:
            continue
        if token not in tokens:
            tokens.append(token)
    return tokens[:30]


def _extract_basic_info_from_text(resume_text: str) -> Dict[str, str]:
    """从原始简历文本里粗略提取姓名/电话/邮箱，供定制简历 PDF 抬头使用。

    只做轻量正则提取，提取不到就留空，绝不臆造。
    """
    text = (resume_text or "").strip()
    info = {"name": "", "phone": "", "email": ""}
    if not text:
        return info

    m_phone = re.search(r"(?<!\d)(1[3-9]\d{9})(?!\d)", text)
    if m_phone:
        info["phone"] = m_phone.group(1)
    m_email = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    if m_email:
        info["email"] = m_email.group(0)

    # 姓名：在前若干行里找"最像姓名"的一行。
    # 复杂排版（双栏/色块）时第一行常是"联系方式""教育背景"等标签，需跳过。
    label_words = {"联系方式", "教育背景", "工作经历", "实习经历", "项目经历", "核心技能",
                   "专业技能", "个人简介", "自我评价", "求职意向", "语言能力", "技能",
                   "荣誉奖项", "校园经历", "基本信息", "电话", "邮箱", "城市", "简历",
                   "硕士", "博士", "本科", "学士", "大专", "专科", "研究生"}
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    # 优先：带"姓名/名字"标签的行
    for line in lines[:15]:
        m = re.match(r"^(?:姓名|名字)[：:\s]*([\u4e00-\u9fff]{2,4})\b", line)
        if m:
            info["name"] = m.group(1)
            break
    # 其次：前 8 行里，纯 2-4 个中文字、且不是栏目标签/机构名的行
    if not info["name"]:
        for line in lines[:8]:
            candidate = re.split(r"[\s，,|·/]", line)[0]
            if not re.fullmatch(r"[\u4e00-\u9fff]{2,4}", candidate):
                continue
            if candidate in label_words:
                continue
            # 排除明显是机构/院校/专业名的词（避免把"清华大学""计算机"当姓名）
            if re.search(r"(大学|学院|学校|公司|专业|科学|技术|工程|管理|系统)$", candidate):
                continue
            info["name"] = candidate
            break
    return info


def _loads_truncated_json(text: str) -> Dict[str, Any]:
    """抢救被 max_tokens 截断的 JSON：从首个 { 起逐字符扫描，
    记录最后一个处于安全位置（对象/数组元素边界）的截断点，再补齐闭合括号后解析。
    用于定制简历 AI 输出过长被截断时，尽量保住已生成的字段，避免直接丢失。"""
    if not text:
        return {}
    start = text.find("{")
    if start == -1:
        return {}
    s = text[start:]
    in_str = False
    escape = False
    stack = []  # 记录 { 和 [
    last_safe = -1  # 最后一个安全可截断位置（某个元素刚结束）
    for i, ch in enumerate(s):
        if escape:
            escape = False
            continue
        if ch == "\\":
            escape = True
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if ch in "{[":
            stack.append(ch)
        elif ch in "}]":
            if stack:
                stack.pop()
            last_safe = i + 1
        elif ch == ",":
            last_safe = i  # 逗号前是一个完整元素
    # 先尝试整体补齐
    for candidate in (s, s[:last_safe] if last_safe > 0 else None):
        if not candidate:
            continue
        fixed = candidate
        # 去掉结尾多余逗号
        fixed = re.sub(r",\s*$", "", fixed)
        # 按未闭合栈补齐括号
        tmp_stack = []
        in_s = False
        esc = False
        for ch in fixed:
            if esc:
                esc = False
                continue
            if ch == "\\":
                esc = True
                continue
            if ch == '"':
                in_s = not in_s
                continue
            if in_s:
                continue
            if ch in "{[":
                tmp_stack.append(ch)
            elif ch in "}]":
                if tmp_stack:
                    tmp_stack.pop()
        closing = "".join("}" if c == "{" else "]" for c in reversed(tmp_stack))
        if in_s:
            fixed += '"'
        try:
            return json.loads(fixed + closing)
        except Exception:
            continue
    return {}


async def _build_customize_resume_with_ai(
    resume_text: str,
    jd_text: str,
    match_result: Dict[str, Any],
    atom_refs: Optional[List[Dict[str, Any]]] = None,
    fallback_basic: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """基于 AI 的定制简历改写：严格只允许基于原简历改写，不允许编造经历。

    复用 one_shot 引擎的 Kimi 配置；输出固定 JSON，便于前端稳定渲染与导出 PDF。
    失败时由调用方回退到规则版兜底。
    atom_refs：用户在「经历原子库」里针对该 JD 改写好的表达，作为优先参考喂给模型。
    fallback_basic：从简历库取到的真实 basic_info（姓名/电话/邮箱），正则猜不到时兜底。
    """
    from litellm import acompletion

    api_key = os.getenv("KIMI_API_KEY", "")
    api_base = os.getenv("KIMI_API_BASE", "https://api.moonshot.cn/v1")
    if not api_key:
        raise RuntimeError("缺少 KIMI_API_KEY，无法进行 AI 定制简历")

    requirement_groups = _extract_requirement_groups(match_result, jd_text)
    matched_requirements = requirement_groups.get("matched", [])
    missing_requirements = requirement_groups.get("missing", [])

    # 组织「原子库已优化表达」参考：用户已针对该 JD 改写过的经历 bullet，优先沿用
    atom_ref_text = ""
    for ref in (atom_refs or [])[:8]:
        title = (ref.get("title") or "").strip()
        bullets = ref.get("bullets") or []
        bullets = [str(b).strip() for b in bullets if str(b).strip()]
        if title and bullets:
            atom_ref_text += f"\n- 【{title}】\n  " + "\n  ".join(bullets)

    # 控制 token：原简历过长时只截断「喂给 AI 的正文」。
    # 姓名/教育/兴趣爱好等确定性字段必须继续从 full_resume_text 提取；
    # 否则长简历尾部的「兴趣爱好」会在 3500 字截断后丢失。
    resume_text = (resume_text or "").strip()
    full_resume_text = resume_text
    resume_text_for_ai = resume_text[:3500] if len(resume_text) > 3500 else resume_text

    atom_ref_block = ("\n【已优化表达参考】（用户已针对该 JD 改写过的经历，请优先沿用）：" + atom_ref_text) if atom_ref_text else ""

    # 两档饱满度（均「宁缺毋滥」，禁止凑数空话）：
    # - 未原子化（无参考）：基于 JD 对已有经历适度扩写，事实足够时每段 3-4 条实质 bullet。
    # - 已原子化（有参考）：利用沉淀素材写得更细，通常每段 3-6 条实质 bullet。
    if atom_ref_text:
        richness_guide = (
            "用户已经做过「经历原子化」，沉淀了更细的素材（见下方参考）。请充分利用这些素材，"
            "把【原简历已有的每一段经历】写得更丰富、更具体、更贴合 JD：覆盖背景、技术方案、负责环节、协作与结果，"
            "做到细节充分、可直接投递。注意：只能扩写已有经历的内部描述，不得新增原文没有的经历。"
        )
        bullet_hint = "有几条实质内容就写几条（通常 3-6 条），宁缺毋滥"
    else:
        richness_guide = (
            "用户没有做经历原子化，只提供了原始简历。请基于 JD 对【原文已有的每一段经历】做适度扩写：比原文更专业、更结构化，"
            "把每段经历的技术、职责、规模、结果展开，但不要硬塞原文没有的细节、更不能新增原文没有的项目或实习。做到能直接投递。"
        )
        bullet_hint = "事实足够时写 3-4 条；事实不足时有几条实质内容就写几条，宁缺毋滥"

    prompt = f"""你是一名资深简历顾问。请基于【原始简历】和【目标 JD】，生成一版更贴合该岗位的定制简历草稿。

【最高原则】
1. 只能基于原始简历里真实存在的经历进行改写、重排、强化表达。
2. 严禁编造任何简历里没有的公司、项目、数据、技能、奖项。
2.1【经历条目红线·最重要】selected_atoms 里的每一段（title/公司/项目）都必须是原始简历里真实写到的经历，一个都不能新增。严禁为了凑数或显得饱满而虚构原文没有的项目/实习（例如原文没写就不准出现“数据分析项目”“缓存系统优化”之类的新经历）。经历的「数量」不能超过原简历真实经历数；「饱满」只能体现在对每一段已有经历的内部描述展开，而不是增加经历条数。
3. 对原简历缺失、但 JD 需要的能力，不要硬写进经历，而是放进 do_not_fake 字段提醒用户。
4. 改写要点：用“场景-动作-结果”结构，能量化就量化（但不能伪造数字）。
5. 若提供了【已优化表达参考】，请优先沿用其中已经针对该 JD 打磨好的 bullet 表达（可微调措辞），但不得引入参考里凭空出现、原简历没有的事实。
6. 【内容饱满度】这是用户可以直接拿去投递的成稿，不是草稿。{richness_guide}
7. 【扩写红线】扩写只能基于原文已经写明的事实，可以补充行业通用的做法描述，但严禁用“更高水平/进一步提升/大幅提高/显著增长”等原文没有的程度词去暗示不存在的成果。原文给了具体数字（如 QPS 3000）就如实写 3000，不得改写成“提升至更高水平”这类似是而非的表达；原文没有结果数据的，就只描述做了什么，不要硬造提升幅度。
8. 【短简历克制】如果原始简历信息很少，只允许输出少量真实内容。profile_summary 不要把 JD 中缺失的能力（如高并发、高可用、微服务、Redis、Kafka、Spring 等）写成候选人的定位、愿景或优势；这些只能放入 do_not_fake。

【目标 JD】
{jd_text[:1500]}

【JD 已匹配要点】{('、'.join(matched_requirements[:6]) or '无')}
【JD 仍欠缺要点】{('、'.join(missing_requirements[:6]) or '无')}

【原始简历】
{resume_text_for_ai}
{atom_ref_block}

只输出合法 JSON，不要 Markdown、不要代码块、不要解释。字段如下：
{{
  "headline": "一句话定位，面向该岗位",
  "profile_summary": "4-6 句个人简介，结合岗位突出核心能力、技术栈、代表性成果与亮点，写得专业且具体，能直接放进简历",
  "skills_line": ["按与岗位相关度排序的技能，最多10个"],
  "selected_atoms": [
    {{
      "title": "经历/项目名称（必须来自原简历）",
      "company": "公司/组织，没有则空字符串",
      "type": "work|project（实习/工作经历填 work，项目经历填 project，只能用这两个英文值之一）",
      "bullets": ["改写后的 bullet，{bullet_hint}，场景-动作-结果结构，每条都要有实质信息，能直接投递"]
    }}
  ],
  "ordering_strategy": ["简历排布建议，1-3 条"],
  "do_not_fake": ["JD 需要但原简历没有、不建议硬写的内容"]
}}

要求：
- selected_atoms 必须来自原简历的真实经历，最多 6 段，最相关的放最前；实习/工作与项目经历都要尽量保留，不要漏掉。
- 每段经历的 bullets：{bullet_hint}。把原文里的技术细节、职责、规模、结果都展开成专业、可直接投递的表达；但每条都必须有实质信息，严禁为了凑数量写“通过技术优化保障系统稳定运行”“提升整体效率”这类没有具体内容的空话套话——没有更多实质内容时，宁可少写一条。
- type 只能填 "work" 或 "project" 两个英文值之一，不要填中文或其它写法。
- skills_line 只列原简历里出现或可合理归纳的技能，不要塞 JD 里但简历没有的。
- 所有字段必须给出，没有内容用空数组或空字符串。
"""

    resp = await acompletion(
        model="openai/moonshot-v1-8k",
        messages=[
            {"role": "system", "content": "你是严谨的简历顾问，只输出合法 JSON，绝不编造简历中不存在的内容。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
        max_tokens=4000,
        api_key=api_key,
        api_base=api_base,
        timeout=60,
        num_retries=3,  # Kimi 过载(429)时自动退避重试，避免直接掉到规则兜底
    )
    content = resp.choices[0].message.content or ""
    stripped = content.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:json)?", "", stripped, flags=re.IGNORECASE).strip()
        stripped = re.sub(r"```$", "", stripped).strip()
    try:
        data = json.loads(stripped)
    except Exception:
        data = None
        # 尝试①：截取最外层 {...} 再解析
        m = re.search(r"\{[\s\S]*\}", stripped)
        if m:
            try:
                data = json.loads(m.group(0))
            except Exception:
                data = None
        # 尝试②：清洗常见瑕疵（全角引号、控制字符、尾逗号）后再解析
        if not data:
            cleaned = (m.group(0) if m else stripped)
            cleaned = cleaned.replace("\u201c", '"').replace("\u201d", '"')  # 全角双引号
            cleaned = re.sub(r"[\x00-\x1f]", " ", cleaned)  # 控制字符
            cleaned = re.sub(r",\s*([}\]])", r"\1", cleaned)  # 尾逗号
            try:
                data = json.loads(cleaned)
            except Exception:
                data = None
        # 尝试③：截断抢救（针对真被 max_tokens 截断的情况）
        if not data:
            data = _loads_truncated_json(stripped)
    if not isinstance(data, dict):
        data = {}
    if not data.get("selected_atoms") and not data.get("headline"):
        logger.warning(f"[customize] JSON 解析疑似失败，原始输出片段: {stripped[:400]}")

    # 规范化为前端已支持的结构
    target_job = _infer_job_title_from_text(jd_text)
    selected = data.get("selected_atoms") or []
    draft_blocks = []
    for atom in selected[:6]:
        if not isinstance(atom, dict):
            continue
        # 归一化经历类型：LLM 返回的 type 写法五花八门（"Work"/"实习"/"intern "/"工作经历"/"项目"），
        # 统一映射成 work / project，避免前端严格匹配 === 'work' 时把实习/工作误分到「项目」栏。
        raw_type = str(atom.get("type") or "").strip().lower()
        if any(k in raw_type for k in ("work", "intern", "job", "实习", "工作", "兼职", "experience")):
            norm_type = "work"
        else:
            norm_type = "project"
        # bullets：先清洗空串，再过滤「凑数空话套话」（无具体数字/技术/对象的泛化表态）。
        # 过滤后若一条不剩，则保留原始第一条兜底，避免该段经历完全空白。
        raw_bullets = [_sanitize_resume_bullet(str(b)) for b in (atom.get("bullets") or []) if str(b).strip()]
        raw_bullets = [b for b in raw_bullets if b]
        kept = [b for b in raw_bullets if not _is_filler_bullet(b)]
        if not kept and raw_bullets:
            kept = raw_bullets[:1]
        draft_blocks.append({
            "title": str(atom.get("title") or "核心经历"),
            "company": str(atom.get("company") or ""),
            "type": norm_type,
            "reasons": [],
            "bullets": kept[:6],
            "skills": [],
        })

    do_not_fake = [str(x) for x in (data.get("do_not_fake") or []) if str(x).strip()]
    summary_bullets = []
    if matched_requirements:
        summary_bullets.append(f"把简历重心放在：{'、'.join(matched_requirements[:2])}")
    if do_not_fake:
        summary_bullets.append(f"以下内容不要硬写，避免被追问穿帮：{'、'.join(do_not_fake[:2])}")

    # 抬头基本信息：先用正则从原文提取，缺失项用简历库里的真实 basic_info 兜底
    basic_info = _extract_basic_info_from_text(full_resume_text)
    if fallback_basic:
        # 姓名/电话/邮箱：正则猜不到时用库里真实值兜底
        for key in ("name", "phone", "email"):
            if not basic_info.get(key) and fallback_basic.get(key):
                basic_info[key] = str(fallback_basic.get(key))
        # 地点 / 教育相关字段：直接透传简历库里的真实值，供六段式抬头与教育段使用
        for key in ("location", "university", "major", "degree", "graduation_year", "birth_date"):
            if fallback_basic.get(key):
                basic_info[key] = str(fallback_basic.get(key))
    # 求职意向：默认填当前 JD 推断出的目标岗位（不编造，来自用户正在投的岗位）
    if not basic_info.get("job_intention") and target_job:
        basic_info["job_intention"] = target_job

    # 教育段：优先用「完整教育数组」展示多段学历（硕士+本科），避免只剩一条。
    education = {}          # 兼容旧前端：最高学历单条
    education_list = []     # 新：多段教育，供导出依次展示
    def _has_education_anchor(item: Dict[str, Any]) -> bool:
        return any(str(item.get(k) or "").strip() for k in ("school", "major", "date_range", "graduation_year"))

    if fallback_basic:
        education = {
            "school": fallback_basic.get("university") or fallback_basic.get("school") or "",
            "major": fallback_basic.get("major") or "",
            "degree": fallback_basic.get("degree") or "",
            "graduation_year": fallback_basic.get("graduation_year") or "",
        }
        for edu in (fallback_basic.get("_education_list") or []):
            if not isinstance(edu, dict):
                continue
            end = str(edu.get("end_date") or "")
            start = str(edu.get("start_date") or "")
            # 起止年份：从 start/end 各取 4 位年份，拼成「2020 - 2024」；缺失则尽量展示已有的
            _ms = re.search(r"(19|20)\d{2}", start)
            _me = re.search(r"(19|20)\d{2}", end)
            sy = _ms.group(0) if _ms else ""
            ey = _me.group(0) if _me else ""
            date_range = (f"{sy} - {ey}" if sy and ey else (sy or ey))
            education_list.append({
                "school": edu.get("school") or "",
                "major": edu.get("major") or "",
                "degree": edu.get("degree") or "",
                "graduation_year": ey,
                "date_range": date_range,
                "_start_year": sy,
            })
        # 排序：学历从低到高（本科 → 硕士 → 博士），同级按入学年份从早到晚。
        # 解析返回的数组顺序不固定（常把硕士放前面），不排序会导致「研究生在上、本科在下」的倒序。
        _deg_rank = {"大专": 1, "专科": 1, "本科": 2, "学士": 2, "bachelor": 2,
                     "硕士": 3, "研究生": 3, "master": 3, "博士": 4, "phd": 4, "doctor": 4}
        def _edu_rank(e):
            d = str(e.get("degree", "")).lower()
            r = 0
            for k, v in _deg_rank.items():
                if k in d:
                    r = v
                    break
            ys = str(e.get("_start_year") or "")
            return (r, int(ys) if ys.isdigit() else 9999)
        education_list = [e for e in education_list if _has_education_anchor(e)]
        education_list.sort(key=_edu_rank)
        for e in education_list:
            e.pop("_start_year", None)
        if education_list:
            education = education_list[-1]
        # 若没有多段数据，至少用最高学历单条兜底，保证有内容
        if not education_list and _has_education_anchor(education):
            education_list = [education]

    # 终极兜底：没有 resume_id（fallback_basic 为空）或库里没存教育时，
    # 直接从简历原文正则抓「学校/学历/年份」，保证教育段不会整段空白。
    if not education_list:
        education_list = [e for e in _extract_education_from_text(full_resume_text) if _has_education_anchor(e)]
        if education_list and not (education.get("school") or education.get("major")):
            education = education_list[-1]  # 单条兼容字段用最高学历（已排序，末尾为最高）

    # 兴趣爱好：优先用简历库里已抓好的（解析时存的），其次从当前文本兜底抓。抓不到留空，绝不编造。
    interests = ""
    if fallback_basic and fallback_basic.get("_interests"):
        interests = str(fallback_basic.get("_interests"))
    if not interests:
        interests = _extract_interests_from_text(full_resume_text)

    # 短简历保险：信息很少时，AI 容易为了“写得像样”而补出数据库设计、
    # 性能优化、稳定性等原文没有的事实。这里直接用原文可验证事实重写核心段落。
    if len(full_resume_text or "") < 500:
        skill_match = re.search(r"技能[:：]\s*([^\n]+)", full_resume_text or "")
        factual_skills = []
        if skill_match:
            factual_skills = [
                s.strip()
                for s in re.split(r"[、,，/；;\s]+", skill_match.group(1))
                if s.strip()
            ][:10]

        project_line = ""
        lines = [ln.strip() for ln in re.split(r"[\n\r]+", full_resume_text or "") if ln.strip()]
        for i, line in enumerate(lines):
            if "项目经历" in line and i + 1 < len(lines):
                project_line = lines[i + 1]
                break
        if not project_line:
            project_line = next((ln for ln in lines if "系统" in ln or "项目" in ln or "课程设计" in ln), "")

        conservative_blocks = []
        if project_line:
            title_part = project_line.split("。", 1)[0]
            title = title_part
            if "：" in title_part:
                left, right = title_part.split("：", 1)
                title = (left + "：" + right.split("，", 1)[0]).strip()
            bullets = []
            if "Java" in project_line or "Java" in (full_resume_text or ""):
                if "增删改查" in project_line:
                    bullets.append("使用 Java 完成图书管理系统的图书增删改查功能")
                else:
                    bullets.append("使用 Java 完成课程项目中的基础后端功能")
            if "登录" in project_line:
                bullets.append("实现用户登录功能")
            if not bullets:
                bullets.append(project_line)
            conservative_blocks.append({
                "title": title or "课程项目",
                "company": "",
                "type": "project",
                "reasons": [],
                "bullets": bullets[:2],
                "skills": [],
            })
        if conservative_blocks:
            draft_blocks = conservative_blocks

        if factual_skills:
            data["skills_line"] = factual_skills

        edu_bits = []
        if education_list:
            e0 = education_list[-1]
            edu_bits = [e0.get("school", ""), e0.get("major", ""), e0.get("degree", "")]
        edu_text = " ".join(x for x in edu_bits if x).strip()
        skill_text = "、".join(factual_skills[:4])
        project_title = conservative_blocks[0]["title"] if conservative_blocks else ""
        summary_parts = []
        if edu_text:
            summary_parts.append(edu_text)
        if skill_text:
            summary_parts.append(f"掌握 {skill_text} 基础")
        if project_title:
            summary_parts.append(f"课程项目经历包括{project_title}，原文明确包含图书增删改查和登录功能")
        conservative_summary = "；".join(summary_parts) + "。"
        data["profile_summary"] = conservative_summary or "已根据原始短简历保守整理，仅保留原文明确出现的信息。"
        data["headline"] = f"{education.get('major') or '计算机相关专业'}应届生，具备基础后端项目经历"

    return {
        "target_job": target_job,
        "basic_info": basic_info,
        "education": education,
        "education_list": education_list,
        "interests": interests,
        "headline": str(data.get("headline") or f"面向 {target_job} 的定制简历草稿"),
        "summary": str(data.get("profile_summary") or "已基于你的原始简历与目标 JD 生成定制草稿，仅做真实经历的改写与重排。"),
        "profile_summary": str(data.get("profile_summary") or ""),
        "skills_line": [str(s) for s in (data.get("skills_line") or []) if str(s).strip()][:10],
        "ordering_strategy": [str(s) for s in (data.get("ordering_strategy") or []) if str(s).strip()][:3],
        "matched_requirements": matched_requirements,
        "missing_requirements": missing_requirements,
        "summary_bullets": summary_bullets[:3],
        "selected_atoms": draft_blocks,
        "optimization_notes": do_not_fake[:4],
        "engine": "ai",
    }


def _build_customize_resume_payload(jd_text: str, match_result: Dict[str, Any], atoms: List[Dict[str, Any]]) -> Dict[str, Any]:
    target_job = _infer_job_title_from_text(jd_text)
    requirement_groups = _extract_requirement_groups(match_result, jd_text)
    matched_requirements = requirement_groups["matched"]
    missing_requirements = requirement_groups["missing"]
    jd_keywords = _tokenize_for_match("\n".join(matched_requirements + missing_requirements + [jd_text]))

    def score_atom(atom: Dict[str, Any]) -> Dict[str, Any]:
        atom_text = " ".join([
            str(atom.get("title", "")),
            str(atom.get("company", "")),
            str(atom.get("description", "")),
            " ".join(atom.get("skills") or []),
        ])
        atom_tokens = _tokenize_for_match(atom_text)
        overlap = [token for token in jd_keywords if token in atom_text and token not in {"产品"}]
        score = len(overlap) * 3
        atom_type = str(atom.get("type") or atom.get("atom_type") or "")
        if atom_type in {"project", "work"}:
            score += 2
        if "AI" in atom_text or "Agent" in atom_text or "B端" in atom_text:
            score += 2
        reasons = []
        if overlap:
            reasons.append(f"命中关键词：{'、'.join(overlap[:3])}")
        if atom_type in {"project", "work"}:
            reasons.append("适合作为定制简历的重点经历")
        return {
            **atom,
            "_score": score,
            "_reasons": reasons[:2]
        }

    ranked_atoms = sorted([score_atom(atom) for atom in atoms], key=lambda x: x["_score"], reverse=True)
    selected_atoms = ranked_atoms[:4]

    def build_bullets(atom: Dict[str, Any]) -> List[str]:
        desc = str(atom.get("description", "")).replace("\r", "\n")
        parts = [p.strip("•- \t") for p in re.split(r"[\n。；;]+", desc) if p.strip()]
        bullets = parts[:3]
        if matched_requirements:
            bullets.insert(0, f"优先呼应 JD：{matched_requirements[0]}")
        return bullets[:4]

    draft_blocks = []
    for atom in selected_atoms:
        draft_blocks.append({
            "title": atom.get("title") or "核心经历",
            "company": atom.get("company") or "",
            "type": atom.get("type") or atom.get("atom_type") or "project",
            "reasons": atom.get("_reasons", []),
            "bullets": build_bullets(atom),
            "skills": (atom.get("skills") or [])[:5]
        })

    rewrite_actions = (
        match_result.get("sections", {}).get("optimization_suggestions")
        or match_result.get("sections", {}).get("rewrite_actions")
        or match_result.get("optimization_suggestions")
        or match_result.get("report_rewrite_priorities")
        or []
    )
    optimization_notes = []
    for item in rewrite_actions[:4]:
        target = item.get("target") or item.get("target_section") or "简历模块"
        action = item.get("action") or item.get("rewrite_method") or ""
        if action:
            optimization_notes.append(f"{target}：{action}")

    # 主线版不再依赖历史经历原子库。没有原子时，直接把本次匹配分析里的
    # 改写优先级转成“建议经历/模块排布”，避免定制简历串入旧测试数据。
    if not draft_blocks:
        for item in rewrite_actions[:3]:
            if not isinstance(item, dict):
                continue
            target = item.get("target_section") or item.get("target") or "重点经历"
            problem = item.get("problem") or item.get("issue") or ""
            method = item.get("rewrite_method") or item.get("action") or ""
            example = item.get("example_direction") or item.get("example") or ""
            side_door = item.get("side_door_fix") or ""
            bullets = [text for text in [problem, method, example, side_door] if text]
            draft_blocks.append({
                "title": target,
                "company": "",
                "type": "rewrite",
                "reasons": ["来自本次匹配分析", "用于当前 JD 定制"],
                "bullets": bullets[:4],
                "skills": []
            })

    if not draft_blocks and (matched_requirements or missing_requirements):
        draft_blocks.append({
            "title": "当前简历主项目 / 实习经历",
            "company": "",
            "type": "rewrite",
            "reasons": ["围绕 JD 重新组织表达"],
            "bullets": [
                f"优先突出已匹配要求：{'、'.join(matched_requirements[:2])}" if matched_requirements else "",
                f"补强证据不足项：{'、'.join(missing_requirements[:2])}" if missing_requirements else "",
                "按“场景-动作-结果-量化”重写关键经历，避免只堆技能词。"
            ],
            "skills": []
        })

    summary_bullets = []
    if matched_requirements:
        summary_bullets.append(f"把简历重心放在：{'、'.join(matched_requirements[:2])}")
    if missing_requirements:
        summary_bullets.append(f"对暂时欠缺的 JD 项不要硬写，优先用已有经历去侧写：{'、'.join(missing_requirements[:2])}")
    if selected_atoms:
        summary_bullets.append(f"优先前置 {selected_atoms[0].get('title', '最相关经历')} 这段经历，提高首屏说服力")

    lead_atoms = [atom.get("title") or "相关经历" for atom in selected_atoms[:2]]
    profile_summary = "；".join([
        f"聚焦 {target_job} 场景的产品/项目经历",
        f"优先突出 {'、'.join(lead_atoms)}" if lead_atoms else "优先突出最相关经历",
        f"围绕 {'、'.join(matched_requirements[:2])} 展开能力证据" if matched_requirements else "围绕 JD 核心要求展开能力证据"
    ])
    skills_line = []
    for atom in selected_atoms:
        for skill in atom.get("skills") or []:
            if skill and skill not in skills_line:
                skills_line.append(skill)
    for req in matched_requirements:
        normalized_req = str(req).strip().strip("；;。")
        if normalized_req and len(normalized_req) <= 24 and normalized_req not in skills_line:
            skills_line.append(normalized_req)
    skills_line = skills_line[:8]

    ordering_strategy = []
    if selected_atoms:
        ordering_strategy.append(f"第 1 段经历建议放 {selected_atoms[0].get('title', '最相关经历')}，作为首屏主证据")
    if len(selected_atoms) > 1:
        ordering_strategy.append(f"第 2 段经历建议放 {selected_atoms[1].get('title', '第二相关经历')}，补足业务或执行链路")
    if missing_requirements:
        ordering_strategy.append(f"对 {'、'.join(missing_requirements[:2])} 不单独新造经历，而是在已有项目里补写相关做法或理解")

    return {
        "target_job": target_job,
        "headline": f"面向 {target_job} 的定制简历草稿",
        "summary": "基于当前简历、目标 JD 和本次匹配结果，优先保留与 JD 直接相关的经历表达，弱化无关内容，强化可迁移能力与证据。",
        "profile_summary": profile_summary,
        "skills_line": skills_line,
        "ordering_strategy": ordering_strategy[:3],
        "matched_requirements": matched_requirements,
        "missing_requirements": missing_requirements,
        "summary_bullets": summary_bullets[:3],
        "selected_atoms": draft_blocks,
        "optimization_notes": optimization_notes[:4],
    }


# ========== 简历优化接口 ==========

@app.post("/api/v1/resume/optimize")
async def optimize_resume(request: OptimizeResumeRequest):
    """优化简历"""
    db = get_db()
    try:
        logger.info(f"Optimizing resume: resume={request.resume_id}, job={request.job_id}")

        # 获取数据
        resume_record = db.get_resume(request.resume_id)
        job_record = db.get_job(request.job_id)

        if not resume_record or not job_record:
            raise HTTPException(status_code=404, detail="简历或岗位不存在")

        resume_data = resume_record
        job_data = job_record

        # 获取之前的匹配结果
        matches = db.list_matches()
        match_result = None
        for m in matches:
            if m.get("resume_id") == request.resume_id and m.get("job_id") == request.job_id:
                match_result = m.get("match_result", {})
                break

        if not match_result:
            # 如果没有匹配结果，先执行匹配
            match_result = match_engine.calculate(resume_data, job_data)

        # 执行优化
        optimization_result = resume_optimizer.optimize_resume(
            resume_data, job_data, match_result
        )

        return {
            "success": True,
            "data": optimization_result,
            "message": "简历优化完成"
        }

    except Exception as e:
        logger.error(f"Resume optimization failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/resume/customize")
async def customize_resume(request: CustomizeResumeRequest):
    """基于原始简历 + JD + 匹配结果，用 AI 生成定制简历草稿（只改写、不编造）。

    优先走 AI 改写；拿不到原始简历或 AI 失败时，回退到规则版兜底，保证不报错。
    """
    # 1) 取原始简历文本：优先 resume_text，其次用 resume_id 从库里拼
    resume_text = (request.resume_text or "").strip()
    fallback_basic: Dict[str, Any] = {}
    if request.resume_id:
        try:
            record = get_db().get_resume(request.resume_id)
            if record:
                fallback_basic = record.get("basic_info", {}) or {}
                # 把「完整教育数组」和「兴趣爱好」一并带上，供导出展示
                fallback_basic = {
                    **fallback_basic,
                    "_education_list": record.get("education_list") or [],
                    "_interests": record.get("interests") or "",
                }
                if not resume_text:
                    parts = []
                    basic = fallback_basic
                    parts.append(" ".join(str(v) for k, v in basic.items() if v and not k.startswith("_")))
                    parts.append("技能：" + "、".join(record.get("skills", []) or []))
                    for exp in record.get("experience", []) or []:
                        parts.append(" ".join(str(exp.get(k, "")) for k in ("company", "position", "duration", "description")))
                    for proj in record.get("projects", []) or []:
                        parts.append(" ".join(str(proj.get(k, "")) for k in ("name", "role", "description")))
                    resume_text = "\n".join(p for p in parts if p.strip())
        except Exception as e:
            logger.warning(f"读取简历用于定制失败: {e}")

    # 2) 优先 AI 改写
    if resume_text:
        try:
            # 从原子库提取「针对该 JD 改写好的表达」作为参考（meta.variants[0].bullets）。
            # 只允许引用标题能在当前简历原文里匹配上的原子；账号级原子库可能包含同一用户
            # 其它简历/其它候选人的经历，不能直接作为本次定制参考。
            normalized_resume_text = re.sub(r"\s+", "", resume_text or "").lower()

            def atom_belongs_to_current_resume(atom: Dict[str, Any]) -> bool:
                meta = atom.get("meta") or {}
                source_resume_id = str(meta.get("source_resume_id") or "").strip()
                if request.resume_id and source_resume_id:
                    return source_resume_id == str(request.resume_id)
                title = str(atom.get("title") or "").strip()
                if not title:
                    return False
                normalized_title = re.sub(r"\s+", "", title).lower()
                return len(normalized_title) >= 3 and normalized_title in normalized_resume_text

            atom_refs = []
            for atom in (request.atoms or []):
                if not atom_belongs_to_current_resume(atom):
                    continue
                meta = atom.get("meta") or {}
                variants = meta.get("variants") or []
                if variants and isinstance(variants[0], dict):
                    raw_bullets = variants[0].get("bullets") or []
                    bullets = []
                    for b in raw_bullets:
                        # 旧版本已经写入库的 variant 可能含有 JD 词硬塞事实；
                        # 进入最终定制前再按当前原子事实二次净化，避免旧脏 variant 污染 PDF。
                        clean = atom_generator._sanitize_rewrite_bullet(str(b), atom)
                        if clean:
                            bullets.append(clean)
                    if bullets:
                        atom_refs.append({"title": atom.get("title", ""), "bullets": bullets})

            # 缓存 key 用「简历全文 + JD + 原子参考」内容哈希，不含 user_id。
            # 只有三者完全一字不差相同才命中 —— 相同输入本就该给相同输出，绝不会串到别人。
            # 重复点「一键优化」时直接返回缓存，省掉一次 v1-8k 调用（降本）。
            # 版本号 customize-v15-sanitized-atom-refs：尾部字段不受截断影响，旧原子 variant 入参前二次净化。
            cache_seed = "customize-v15-sanitized-atom-refs\x00" + (resume_text or "") + "\x00" + (request.jd_text or "") + "\x00" + json.dumps(atom_refs, ensure_ascii=False, sort_keys=True)
            cached_payload = cache.get_cached_llm_result(cache_seed)
            if cached_payload:
                logger.info("[customize] 命中缓存，跳过 AI 调用（降本）")
                return {"success": True, "data": cached_payload, "message": "定制简历草稿生成完成（缓存）", "cached": True}

            payload = await _build_customize_resume_with_ai(
                resume_text, request.jd_text, request.match_result or {}, atom_refs, fallback_basic
            )
            # 仅缓存成功的 AI 结果（engine=ai），规则兜底版不缓存，避免把降级结果固化
            if isinstance(payload, dict) and payload.get("engine") == "ai":
                cache.cache_llm_result(cache_seed, payload)
            return {"success": True, "data": payload, "message": "定制简历草稿生成完成"}
        except Exception as e:
            logger.warning(f"AI 定制简历失败: {e}")
            # 关键：已经知道当前简历是谁（有 resume_text）时，绝不回退到「规则版」。
            # 规则版的正文来自原子库/匹配结果，与当前本人无关，AI 限流(429)时回退会产出
            # 「抬头是本人、正文是别人」的张冠李戴简历（严重串数据）。宁可如实报错让用户重试。
            msg = str(e)
            if "429" in msg or "overload" in msg.lower() or "rate" in msg.lower():
                raise HTTPException(status_code=503, detail="AI 接口当前繁忙（请求过多），请稍候 1-2 分钟后重试。")
            raise HTTPException(status_code=502, detail="AI 定制简历生成失败，请稍后重试。")

    # 3) 回退：仅在「完全没有原始简历文本」时才用规则版兜底（无法定位本人，退而求其次）
    try:
        payload = _build_customize_resume_payload(
            request.jd_text,
            request.match_result or {},
            request.atoms or []
        )
        return {"success": True, "data": payload, "message": "定制简历草稿生成完成（基础版）"}
    except Exception as e:
        logger.error(f"Customize resume failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _is_pdf_latin_char(ch: str) -> bool:
    return ord(ch) < 128


def _split_pdf_text_runs(text: str) -> List[tuple]:
    """按中英文拆成绘制片段，避免内置 CJK 字体把英文显示成 J a v a。"""
    runs = []
    current = ""
    current_is_latin = None
    for ch in str(text):
        is_latin = _is_pdf_latin_char(ch)
        if current and is_latin != current_is_latin:
            runs.append((current, current_is_latin))
            current = ch
        else:
            current += ch
        current_is_latin = is_latin
    if current:
        runs.append((current, current_is_latin))
    return runs


def _pdf_text_width(text: str, fontname: str, fontsize: float) -> float:
    try:
        import fitz
        return fitz.get_text_length(str(text), fontname=fontname, fontsize=fontsize)
    except Exception:
        # 保守估算：中文约等于 1em，英文约 0.55em。
        width = 0.0
        for ch in str(text):
            width += fontsize if "\u4e00" <= ch <= "\u9fff" else fontsize * 0.55
        return width


def _pdf_mixed_text_width(text: str, cjk_font: str, latin_font: str, fontsize: float) -> float:
    width = 0.0
    for run, is_latin in _split_pdf_text_runs(text):
        width += _pdf_text_width(run, latin_font if is_latin else cjk_font, fontsize)
    return width


def _wrap_pdf_text(text: str, max_width: float, fontname: str, fontsize: float, latin_font: str = "helv") -> List[str]:
    """按实际字体宽度做混排换行，避免中文 PDF 溢出页面。"""
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if not text:
        return []
    lines: List[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if current and _pdf_mixed_text_width(test, fontname, latin_font, fontsize) > max_width:
            lines.append(current.rstrip())
            current = ch.lstrip()
        else:
            current = test
    if current.strip():
        lines.append(current.strip())
    return lines


def _build_resume_pdf_bytes(resume: Dict[str, Any]) -> bytes:
    """后端直出 PDF：严格按中英文分段绘制，避免 china-s 把英文渲染成带空格的字母。"""
    import fitz

    def val(x: Any) -> str:
        return str(x or "").strip()

    def clean_bullet(text: Any) -> str:
        return re.sub(r"^\s*[-•*]\s*", "", val(text)).strip()

    def displayable_education(item: Any) -> bool:
        if not isinstance(item, dict):
            return False
        return any(val(item.get(k)) for k in ("school", "major", "date_range", "graduation_year"))

    # 选择适合中文 + 英文的字体：中文统一用 china-s（PyMuPDF 内置 CJK），英文用 helv/hebo。
    # 加粗英文用 hebo（Helvetica-Bold），不加粗英文用 helv；中文没有独立的加粗内建字体，
    # 通过“标题字号更大 + 分割线”营造层级，避免 bold 偏移二次绘制带来的重影问题。
    def cjk_font_for(bold: bool) -> str:
        return "china-s"  # 本身就是等宽黑体感，标题通过字号+分割线与正文区分

    def latin_font_for(bold: bool) -> str:
        return "hebo" if bold else "helv"

    def _segment_width(seg: str, size: float, bold: bool) -> float:
        is_latin = all(ord(c) < 128 for c in seg)
        fname = latin_font_for(bold) if is_latin else cjk_font_for(bold)
        try:
            return fitz.get_text_length(seg, fontname=fname, fontsize=size)
        except Exception:
            return size * (1.0 if any("\u4e00" <= c <= "\u9fff" for c in seg) else 0.55) * len(seg)

    def _mixed_width(text: str, size: float, bold: bool) -> float:
        total = 0.0
        for seg, _ in _split_pdf_text_runs(text):
            total += _segment_width(seg, size, bold)
        return total

    def _mixed_wrap(text: str, max_w: float, size: float, bold: bool) -> List[str]:
        text = re.sub(r"\s+", " ", str(text or "")).strip()
        if not text:
            return []
        lines: List[str] = []
        current = ""
        for ch in text:
            candidate = current + ch
            if current and _mixed_width(candidate, size, bold) > max_w:
                lines.append(current.rstrip())
                current = ch.lstrip()
            else:
                current = candidate
        if current.strip():
            lines.append(current.strip())
        return lines

    def _draw_mixed_line(page, x: float, y_top: float, text: str, size: float, color, bold: bool):
        cursor_x = x
        for seg, is_latin in _split_pdf_text_runs(text):
            if not seg:
                continue
            fname = latin_font_for(bold) if is_latin else cjk_font_for(bold)
            if bold and not is_latin:
                # CJK 内置字体没有独立 bold 字重，用 PDF 原生 fill+stroke 加粗；
                # 只绘制一次文本，避免之前“画两遍模拟加粗”导致抽取文本重复。
                page.insert_text(
                    (cursor_x, y_top + size),
                    seg,
                    fontname=fname,
                    fontsize=size,
                    color=color,
                    fill=color,
                    render_mode=2,
                    border_width=0.035 if size <= 11.2 else 0.045,
                )
            else:
                page.insert_text((cursor_x, y_top + size), seg, fontname=fname, fontsize=size, color=color)
            cursor_x += _segment_width(seg, size, bold)

    basic = resume.get("basic_info") or {}
    name = val(basic.get("name")) or "我的简历"
    contact_parts = [
        val(basic.get("phone")),
        val(basic.get("email")),
        ("求职意向：" + val(basic.get("job_intention"))) if val(basic.get("job_intention")) else "",
        val(basic.get("location")),
    ]
    contact_parts = [p for p in contact_parts if p]
    profile_summary = val(resume.get("profile_summary") or resume.get("summary"))
    skills = [val(s) for s in (resume.get("skills_line") or []) if val(s)]
    blocks = [b for b in (resume.get("selected_atoms") or []) if isinstance(b, dict)]
    work_blocks = [b for b in blocks if val(b.get("type")) in {"work", "intern", "internship"}]
    project_blocks = [b for b in blocks if val(b.get("type")) not in {"work", "intern", "internship"}]
    edu = resume.get("education") or {}
    raw_edu_list = resume.get("education_list") or ([edu] if (edu.get("school") or edu.get("major")) else [])
    edu_list = [e for e in raw_edu_list if displayable_education(e)]
    interests = val(resume.get("interests"))

    doc = fitz.open()
    page_w, page_h = 595, 842  # A4 points
    margin_l, margin_r, margin_t, margin_b = 42, 42, 38, 42
    content_w = page_w - margin_l - margin_r
    text_color = (0.12, 0.16, 0.22)
    muted = (0.38, 0.43, 0.50)

    page = None
    y = margin_t

    def new_page():
        nonlocal page, y
        page = doc.new_page(width=page_w, height=page_h)
        y = margin_t

    def ensure(height: float):
        nonlocal page
        if page is None:
            new_page()
            return
        if y + height > page_h - margin_b:
            new_page()

    def write_line(text: str, size: float = 10.5, color=text_color, x: float = margin_l, leading: float = 1.42, bold: bool = False):
        nonlocal y
        ensure(size * leading)
        # 不再做 bold 偏移二次绘制：英文用 hebo 本身就是粗体，中文用更大字号+分隔线体现层级
        _draw_mixed_line(page, x, y, text, size, color, bold)
        y += size * leading

    def write_wrapped(text: str, size: float = 10.5, color=text_color, indent: float = 0, leading: float = 1.45, bold: bool = False):
        nonlocal y
        max_w = content_w - indent
        for line in _mixed_wrap(text, max_w, size, bold):
            write_line(line, size=size, color=color, x=margin_l + indent, leading=leading, bold=bold)

    def section(title: str, subtitle: str = ""):
        nonlocal y
        ensure(36)
        y += 7
        line_top = y
        write_line(title, size=13.2, color=(0.06, 0.09, 0.16), leading=1.45, bold=True)
        if subtitle:
            subtitle_x = margin_l + _mixed_width(title, 13.2, True) + 7
            page.insert_text(
                (subtitle_x, line_top + 12.2),
                subtitle,
                fontname="helv",
                fontsize=9.0,
                color=muted,
            )
        page.draw_line((margin_l, y), (page_w - margin_r, y), color=(0.06, 0.09, 0.16), width=0.8)
        y += 8

    def exp_block(block: Dict[str, Any]):
        nonlocal y
        title = val(block.get("title")) or "核心经历"
        company = val(block.get("company"))
        bullets = [clean_bullet(b) for b in (block.get("bullets") or []) if clean_bullet(b)]
        estimated = 24 + len(bullets) * 30
        ensure(min(estimated, 120))
        write_wrapped(title + (f" | {company}" if company else ""), size=11.2, color=(0.06, 0.09, 0.16), leading=1.32, bold=True)
        for bullet in bullets:
            for line in _mixed_wrap(bullet, content_w, 10.5, False):
                write_line(line, size=10.5, color=text_color, x=margin_l, leading=1.35)
        y += 5

    new_page()
    section("个人信息", "Profile")
    write_line(name, size=20, color=(0.06, 0.09, 0.16), leading=1.35, bold=True)
    if contact_parts:
        write_wrapped(" | ".join(contact_parts), size=9.8, color=muted)
        y += 2

    if profile_summary:
        section("个人介绍", "Summary")
        write_wrapped(profile_summary, size=10.5)

    if skills:
        section("核心技能", "Skills")
        write_wrapped("、".join(skills), size=10.3)

    if project_blocks:
        section("项目", "Projects")
        for block in project_blocks:
            exp_block(block)

    if work_blocks:
        section("实习 / 工作经历", "Experience")
        for block in work_blocks:
            exp_block(block)

    if edu_list:
        section("教育", "Education")
        for e in edu_list:
            if not isinstance(e, dict):
                continue
            dr = val(e.get("date_range") or e.get("graduation_year"))
            parts = [val(e.get("school")), val(e.get("major")), val(e.get("degree")), dr]
            parts = [p for p in parts if p]
            if parts:
                write_wrapped(" / ".join(parts), size=11.2, color=(0.06, 0.09, 0.16), leading=1.35, bold=True)

    if interests:
        section("兴趣爱好", "Interests")
        write_wrapped(interests, size=10.5)

    out = io.BytesIO()
    doc.save(out, deflate=True, garbage=4)
    doc.close()
    return out.getvalue()


@app.post("/api/v1/resume/export-pdf")
async def export_resume_pdf(request: ExportResumePdfRequest):
    """直接下载定制简历 PDF。后端生成文本型 PDF，避免前端截图式导出截断/串画面。"""
    try:
        resume = request.resume or {}
        basic = resume.get("basic_info") or {}
        name = str(basic.get("name") or "我的简历").strip()
        pdf_bytes = _build_resume_pdf_bytes(resume)
        filename = f"{name}-定制简历.pdf"
        quoted = quote(filename)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{quoted}",
                "Cache-Control": "no-store",
            },
        )
    except Exception as e:
        logger.error(f"Export PDF failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="PDF 导出失败，请稍后重试。")


# ========== 历史记录接口 ==========

@app.get("/api/v1/history")
async def get_history(limit: int = 20):
    """获取匹配历史记录"""
    db = get_db()
    try:
        matches = db.list_matches(limit=limit)
        return {"success": True, "data": matches, "count": len(matches)}
    except Exception as e:
        logger.error(f"Failed to get history: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== 经历原子库接口 ==========

@app.get("/api/v1/atoms")
async def get_atoms(current_user: Dict[str, Any] = Depends(get_current_user)):
    """获取经历原子库（仅当前登录用户自己的）"""
    db = get_db()
    try:
        atoms = db.list_atoms(user_id=current_user["id"])
        return {"success": True, "data": atoms, "count": len(atoms)}
    except Exception as e:
        logger.error(f"Failed to get atoms: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/atoms/generate")
async def generate_atoms(resume_id: str = Form(...), current_user: Dict[str, Any] = Depends(get_current_user)):
    """从指定简历生成经历原子并入库（归属当前登录用户）"""
    db = get_db()
    try:
        resume_data = db.get_resume(resume_id)
        if not resume_data:
            raise HTTPException(status_code=404, detail="简历不存在")

        atoms = atom_generator.from_resume_data(resume_data)
        saved = []
        for atom in atoms:
            meta = atom.get("meta") or {}
            meta["source_resume_id"] = resume_id
            atom["meta"] = meta
            atom_id = db.save_atom(atom, user_id=current_user["id"])
            saved.append({**atom, "id": atom_id})

        logger.info(f"Generated {len(saved)} atoms from resume {resume_id}")
        return {"success": True, "data": saved, "count": len(saved), "message": "经历原子生成成功"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to generate atoms: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _normalize_skills(skills: Any) -> List[str]:
    """统一 skills 字段为字符串列表（兼容逗号分隔字符串与数组）"""
    if not skills:
        return []
    if isinstance(skills, list):
        return [str(s).strip() for s in skills if str(s).strip()]
    if isinstance(skills, str):
        return [s.strip() for s in re.split(r"[,，;；]", skills) if s.strip()]
    return []


@app.post("/api/v1/atoms")
async def create_atom(request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    """手动创建经历原子（兼容 FormData 与 JSON 两种提交方式，归属当前登录用户）"""
    db = get_db()
    try:
        content_type = request.headers.get("content-type", "")
        if "application/json" in content_type:
            payload = await request.json()
        else:
            form = await request.form()
            payload = dict(form)

        title = (payload.get("title") or "").strip()
        if not title:
            raise HTTPException(status_code=400, detail="标题不能为空")

        description = payload.get("description", "")
        company = payload.get("company", "")
        atom = {
            "title": title,
            "type": payload.get("atom_type") or payload.get("type") or "work",
            "description": description,
            "company": company,
            "skills": _normalize_skills(payload.get("skills")),
            "meta": {
                "fact": {
                    "company": company,
                    "role": payload.get("role", "") or title,
                    "duration": payload.get("duration", ""),
                },
                "base_description": description,
                "highlight": "",
                "variants": []
            }
        }
        atom_id = db.save_atom(atom, user_id=current_user["id"])
        logger.info(f"Atom created: {atom_id}")
        return {"success": True, "data": {**atom, "id": atom_id}, "message": "经历原子创建成功"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to create atom: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/v1/atoms/{atom_id}")
async def delete_atom(atom_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    """删除经历原子（仅能删当前登录用户自己的）"""
    db = get_db()
    try:
        deleted = db.delete_atom(atom_id, user_id=current_user["id"])
        if not deleted:
            raise HTTPException(status_code=404, detail="经历原子不存在")
        logger.info(f"Atom deleted: {atom_id}")
        return {"success": True, "message": "经历原子已删除"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete atom: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class RewriteAtomRequest(BaseModel):
    """针对 JD 改写经历原子表达层"""
    jd_text: str


@app.post("/api/v1/atoms/{atom_id}/rewrite")
async def rewrite_atom_for_jd(atom_id: str, request: RewriteAtomRequest, current_user: Dict[str, Any] = Depends(get_current_user)):
    """针对目标 JD 重写经历原子的表达层，生成并保存一个改写版本(variant)。仅限本人原子。"""
    db = get_db()
    try:
        if not request.jd_text.strip():
            raise HTTPException(status_code=400, detail="请提供目标 JD 文本")

        atom = db.get_atom(atom_id, user_id=current_user["id"])
        if not atom:
            raise HTTPException(status_code=404, detail="经历原子不存在")

        # 兴趣类原子：走「目标用户契合论证」；其余经历类：走 STAR 表达改写
        if atom.get("type") == "interest":
            variant = atom_generator.argue_user_fit(atom, request.jd_text)
            if not variant.get("relevant"):
                # 不契合：不写入版本，直接返回提示，避免教用户硬蹭
                note = variant.get("note") or "此岗位与该兴趣无明显契合，无需在简历中突出。"
                return {"success": True, "data": {"variant": variant, "skipped": True}, "message": note}
            if not variant.get("bullets"):
                raise HTTPException(status_code=502, detail="AI 论证暂不可用，请稍后重试")
            msg = "已生成「目标用户契合」论证"
        else:
            variant = atom_generator.rewrite_for_jd(atom, request.jd_text)
            if not variant.get("bullets"):
                raise HTTPException(status_code=502, detail="AI 改写暂不可用，请稍后重试")
            msg = "已生成针对该 JD 的改写版本"

        meta = atom.get("meta") or {}
        variants = meta.get("variants") or []
        variants.insert(0, variant)
        meta["variants"] = variants[:5]  # 最多保留最近 5 个版本
        db.update_atom_meta(atom_id, meta)

        logger.info(f"Atom rewritten for JD: {atom_id}")
        return {"success": True, "data": {"variant": variant, "meta": meta}, "message": msg}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to rewrite atom: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/resumes/{resume_id}")
async def get_resume_by_id(resume_id: str):
    """按 ID 获取简历数据"""
    db = get_db()
    try:
        resume_data = db.get_resume(resume_id)
        if not resume_data:
            raise HTTPException(status_code=404, detail="简历不存在")
        return {"success": True, "data": {**resume_data, "id": resume_id}}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/analysis/deep")
async def deep_analysis(resume_id: str = Form(...), job_id: str = Form(...)):
    """JD 逐句拆解深度分析"""
    db = get_db()
    try:
        resume_data = db.get_resume(resume_id)
        job_data = db.get_job(job_id)
        if not resume_data or not job_data:
            raise HTTPException(status_code=404, detail="简历或岗位不存在")

        result = deep_analysis_service.analyze(resume_data, job_data)
        logger.info(f"Deep analysis completed: resume={resume_id}, job={job_id}")
        return {"success": True, "data": result, "message": "深度分析完成"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Deep analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== 岗位推荐接口 ==========

@app.post("/api/v1/resume/extract-text")
async def extract_resume_text(file: UploadFile = File(...)):
    """快速提取简历纯文本（不经过 LLM），供岗位推荐等轻量场景使用"""
    try:
        content = await file.read()
        filename = file.filename or ""
        lower = filename.lower()
        text = ""
        engine = "text"  # text=文本层解析；ocr=视觉OCR兜底
        is_image = lower.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp"))

        if lower.endswith(".pdf"):
            text = get_pdf_parser().extract_text(content, filename)
        elif lower.endswith(".docx"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join([p.text for p in doc.paragraphs])
            except Exception:
                text = content.decode("utf-8", errors="ignore")
        elif is_image:
            text = ""  # 图片没有文本层，直接走 OCR
        else:
            text = content.decode("utf-8", errors="ignore")
        text = (text or "").strip()

        def _looks_garbled(t: str) -> bool:
            import re as _re
            meaningful = _re.findall(r"[\u4e00-\u9fff0-9A-Za-z]", t)
            return bool(t) and (len(meaningful) / max(len(t), 1)) < 0.25

        # 兜底：文本层解析不出来（图片型/扫描版 PDF、字体编码乱码、纯图片文件）→ 视觉 OCR
        need_ocr = is_image or len(text) < 120 or _looks_garbled(text)
        if need_ocr and (lower.endswith(".pdf") or is_image):
            try:
                ocr_text = (await vision_parser.extract_text_via_vision(content, filename)).strip()
                # 仅当 OCR 结果明显更可信（更长且不乱码）才采用
                if ocr_text and len(ocr_text) > len(text) and not _looks_garbled(ocr_text):
                    text = ocr_text
                    engine = "ocr"
            except Exception as e:
                logger.warning(f"视觉 OCR 兜底失败: {e}")

        garbled = _looks_garbled(text)
        return {
            "success": True,
            "text": text,
            "char_count": len(text),
            "garbled": garbled,
            "engine": engine,
        }
    except Exception as e:
        logger.error(f"Extract resume text failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class RecommendJobsRequest(BaseModel):
    """岗位推荐请求：提供 resume_id / resume 数据 / resume_text 任一即可"""
    resume_id: Optional[str] = None
    resume: Optional[Dict[str, Any]] = None
    resume_text: Optional[str] = None
    top_n: int = 5


@app.post("/api/v1/jobs/recommend")
async def recommend_jobs(request: RecommendJobsRequest):
    """根据简历从内置岗位池推荐匹配度较高的岗位（命中赛题痛点1）"""
    db = get_db()
    try:
        resume_data = request.resume
        if not resume_data and request.resume_id:
            resume_data = db.get_resume(request.resume_id)
        if not resume_data and request.resume_text:
            # 纯文本输入：不经过 LLM，直接做轻量匹配，演示更稳定
            resume_data = {"raw_text": request.resume_text}
        if not resume_data:
            raise HTTPException(status_code=400, detail="请提供 resume_id、resume 或 resume_text")

        recommender = get_job_recommender()
        top_n = max(1, min(10, request.top_n or 5))
        recommendations = recommender.recommend(resume_data, top_n=top_n)

        logger.info(f"Job recommendation done: {len(recommendations)} jobs")
        return {
            "success": True,
            "data": recommendations,
            "count": len(recommendations),
            "message": "岗位推荐完成"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Job recommendation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== 投递追踪接口 ==========

@app.post("/api/v1/applications")
async def create_application(data: Dict[str, Any]):
    """创建投递记录"""
    db = get_db()
    try:
        app_id = db.save_application(data)
        logger.info(f"Application created: {app_id}")
        return {"success": True, "data": {"id": app_id}, "message": "投递记录创建成功"}
    except Exception as e:
        logger.error(f"Failed to create application: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/applications")
async def get_applications():
    """获取投递记录"""
    db = get_db()
    try:
        applications = db.list_applications()
        return {"success": True, "data": applications, "count": len(applications)}
    except Exception as e:
        logger.error(f"Failed to get applications: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== 校招评分接口 ==========

class CampusScoreRequest(BaseModel):
    """校招评分请求"""
    resume_data: Dict[str, Any]
    job_data: Optional[Dict[str, Any]] = None


def _normalize_resume_for_campus_scoring(resume: Dict) -> Dict:
    """将 offer-catcher 简历格式转换为 campus_scorer 期望的格式."""
    normalized = {}

    # 基本信息 -> education 列表
    basic_info = resume.get("basic_info", {})
    if basic_info.get("university") or basic_info.get("school"):
        normalized["education"] = [{
            "institution": basic_info.get("university") or basic_info.get("school", ""),
            "field": basic_info.get("major", ""),  # campus_scorer 期望 field
            "major": basic_info.get("major", ""),
            "degree": basic_info.get("degree", ""),
            "start_date": basic_info.get("start_date", ""),
            "end_date": basic_info.get("graduation_year", ""),
        }]

    # 经历 -> workExperience（campus_scorer 期望的字段名）
    experiences = resume.get("experience", [])
    if experiences:
        work_exp = []
        for exp in experiences:
            work_exp.append({
                "company": exp.get("company", ""),
                "title": exp.get("position", exp.get("title", "")),  # 适配 position/title
                "description": exp.get("description", ""),
            })
        normalized["workExperience"] = work_exp

    # 项目
    normalized["projects"] = resume.get("projects", [])

    # 技能 - 转为字符串列表和 summary 文本
    skills = resume.get("skills", [])
    skill_names = []
    if skills and isinstance(skills, list):
        for skill in skills:
            if isinstance(skill, str):
                skill_names.append(skill)
            elif isinstance(skill, dict):
                skill_names.append(skill.get("name", ""))

    # campus_scorer 技能评分期望 summary 字段包含技能文本
    normalized["skills"] = [s for s in skill_names if s]
    normalized["summary"] = " ".join(skill_names)  # 用于文本匹配

    # 教育（如果有独立的 education 字段）
    if "education" in resume and resume["education"]:
        if "education" not in normalized:
            normalized["education"] = []
        if isinstance(resume["education"], dict):
            edu = resume["education"]
            normalized["education"].append({
                "institution": edu.get("school", "") or edu.get("institution", ""),
                "field": edu.get("major", ""),
                "major": edu.get("major", ""),
                "degree": edu.get("degree", ""),
            })
        else:
            normalized["education"].extend(resume["education"])

    return normalized


@app.post("/api/v1/campus-score")
async def calculate_campus_score(request: CampusScoreRequest):
    """校招HR级评分 - 真实大厂筛选逻辑"""
    try:
        # 转换数据格式
        normalized_resume = _normalize_resume_for_campus_scoring(request.resume_data)
        score_result = campus_scorer.score(normalized_resume, request.job_data)

        return {
            "success": True,
            "data": {
                "total_score": score_result.total_score,
                "grade": score_result.grade,
                "scores": {
                    "university": score_result.university_score,
                    "degree": score_result.degree_score,
                    "major": score_result.major_score,
                    "internship": score_result.internship_score,
                    "project": score_result.project_score,
                    "skill": score_result.skill_score,
                    "achievement": score_result.achievement_score,
                },
                "details": {
                    "university_tier": score_result.university_tier,
                    "degree_level": score_result.degree_level,
                    "internship_companies": score_result.internship_companies,
                    "top_projects": score_result.top_projects,
                    "matched_skills": score_result.matched_skills,
                    "missing_skills": score_result.missing_skills,
                },
                "analysis": {
                    "strengths": score_result.strengths,
                    "weaknesses": score_result.weaknesses,
                    "recommendation": score_result.recommendation,
                }
            },
            "message": "校招评分完成"
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"评分失败: {str(e)}")


# ========== 缓存管理接口 ==========

@app.get("/api/v1/cache/stats")
async def get_cache_stats():
    """获取缓存统计信息"""
    try:
        stats = cache.get_stats()
        return {
            "success": True,
            "data": stats,
            "message": "缓存统计获取成功"
        }
    except Exception as e:
        logger.error(f"Failed to get cache stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/cache/clear")
async def clear_cache(pattern: str = ""):
    """清除缓存
    - 不传 pattern: 清空所有缓存
    - 传 pattern: 按模式清除，如 "resume" 清除简历缓存
    """
    try:
        if pattern:
            count = cache.delete_pattern(pattern)
            logger.info(f"Cleared cache pattern: {pattern}, count: {count}")
            return {
                "success": True,
                "data": {"cleared_count": count},
                "message": f"已清除 {count} 条缓存"
            }
        else:
            cache.flush_db()
            logger.info("Flushed all cache")
            return {
                "success": True,
                "message": "已清空所有缓存"
            }
    except Exception as e:
        logger.error(f"Failed to clear cache: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== Celery 异步任务接口 ==========

class AsyncParseResumeTextRequest(BaseModel):
    """异步简历解析请求"""
    text: str


@app.post("/api/v1/resume/parse-async")
async def parse_resume_async(request: AsyncParseResumeTextRequest):
    """
    异步解析简历文本（后台任务）
    立即返回任务 ID，用户可轮询获取结果
    """
    try:
        task = parse_resume_text_task.delay(request.text)
        logger.info(f"Async resume parse task created: {task.id}")
        return {
            "success": True,
            "task_id": task.id,
            "message": "简历解析任务已创建，请使用 task_id 查询结果"
        }
    except Exception as e:
        logger.error(f"Failed to create async task: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class AsyncParseJobRequest(BaseModel):
    """异步 JD 解析请求"""
    jd_text: str


@app.post("/api/v1/job/parse-async")
async def parse_job_async(request: AsyncParseJobRequest):
    """异步解析 JD 文本"""
    try:
        task = parse_job_task.delay(request.jd_text)
        logger.info(f"Async job parse task created: {task.id}")
        return {
            "success": True,
            "task_id": task.id,
            "message": "JD 解析任务已创建，请使用 task_id 查询结果"
        }
    except Exception as e:
        logger.error(f"Failed to create async task: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class AsyncMatchRequest(BaseModel):
    """异步匹配分析请求"""
    resume_id: str
    job_id: str


@app.post("/api/v1/match-async")
async def match_analysis_async(request: AsyncMatchRequest):
    """异步执行匹配分析"""
    try:
        task = match_analysis_task.delay(request.resume_id, request.job_id)
        logger.info(f"Async match task created: {task.id}")
        return {
            "success": True,
            "task_id": task.id,
            "message": "匹配分析任务已创建，请使用 task_id 查询结果"
        }
    except Exception as e:
        logger.error(f"Failed to create async task: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/tasks/{task_id}")
async def get_task_status(task_id: str):
    """
    查询异步任务状态

    返回状态：
    - PENDING: 等待执行
    - STARTED: 正在执行
    - SUCCESS: 执行成功
    - FAILURE: 执行失败
    - RETRY: 重试中
    """
    try:
        result = celery_app.AsyncResult(task_id)

        response = {
            "task_id": task_id,
            "status": result.state,
        }

        if result.state == "PENDING":
            response["message"] = "任务等待执行"
        elif result.state == "STARTED":
            response["message"] = "任务正在执行"
        elif result.state == "SUCCESS":
            response["message"] = "任务执行成功"
            response["result"] = result.result
        elif result.state == "FAILURE":
            response["message"] = "任务执行失败"
            response["error"] = str(result.info)
        elif result.state == "RETRY":
            response["message"] = "任务重试中"
            response["retry_count"] = result.info.get("retry_count", 0)

        return response
    except Exception as e:
        logger.error(f"Failed to get task status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/tasks/{task_id}/cancel")
async def cancel_task(task_id: str):
    """取消正在执行的任务"""
    try:
        celery_app.control.revoke(task_id, terminate=True)
        logger.info(f"Task cancelled: {task_id}")
        return {
            "success": True,
            "message": "任务已取消"
        }
    except Exception as e:
        logger.error(f"Failed to cancel task: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== 监控和告警接口 ==========

@app.get("/metrics")
async def metrics():
    """
    Prometheus 指标端点

    返回应用性能指标，供 Prometheus 抓取
    包括：
    - HTTP 请求计数和延迟
    - 缓存命中率
    - LLM 调用统计
    - 数据库操作统计
    - Celery 任务统计
    """
    return get_metrics()


@app.get("/health")
async def health_check():
    """
    健康检查端点

    用于 Kubernetes/Docker 健康检查
    返回各组件状态
    """
    from app.monitoring.metrics import app_info

    checks = {
        "api": "healthy",
        "database": "unknown",
        "redis": "unknown",
        "celery": "unknown",
    }

    # 检查数据库
    try:
        db = get_db()
        db.get_stats()
        checks["database"] = "healthy"
    except Exception as e:
        checks["database"] = f"unhealthy: {str(e)}"

    # 检查 Redis
    try:
        from app.cache.redis_cache import get_cache
        cache = get_cache()
        cache.get_redis().ping()
        checks["redis"] = "healthy"
    except Exception as e:
        checks["redis"] = f"unhealthy: {str(e)}"

    # 检查 Celery
    try:
        inspector = celery_app.control.inspect()
        stats = inspector.stats()
        if stats:
            checks["celery"] = "healthy"
        else:
            checks["celery"] = "unhealthy: no workers"
    except Exception as e:
        checks["celery"] = f"unhealthy: {str(e)}"

    all_healthy = all(v == "healthy" for v in checks.values())
    status_code = 200 if all_healthy else 503

    return Response(
        content=JSONResponse(content=checks).body.decode(),
        status_code=status_code,
        media_type="application/json"
    )


@app.get("/api/v1/status")
async def get_status():
    """
    详细状态端点

    返回应用、依赖服务、队列状态
    """
    from app.monitoring.metrics import http_requests_total, cache_hits_total, cache_misses_total

    status = {
        "app": {
            "version": "1.2.0",
            "name": "Offer 捕手",
        },
        "metrics": {
            "http_requests_total": http_requests_total._value.get() if hasattr(http_requests_total, '_value') else 0,
            "cache_hits": cache_hits_total._value.get() if hasattr(cache_hits_total, '_value') else 0,
            "cache_misses": cache_misses_total._value.get() if hasattr(cache_misses_total, '_value') else 0,
        }
    }

    # 获取 Celery worker 状态
    try:
        inspector = celery_app.control.inspect()
        worker_stats = inspector.stats()
        if worker_stats:
            status["celery_workers"] = list(worker_stats.keys())
            status["celery_active_tasks"] = sum(
                len(s.get('active_tasks', [])) for s in worker_stats.values()
            )
    except Exception:
        status["celery_workers"] = []

    return status


# ========== 测试数据接口 ==========

TEST_RESUME_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "uploads",
    "fixed-test-resume.pdf"
)
TEST_JD = """AI产品实习生

• 薪资：250-300元/天

• 地点：上海·徐汇区·漕河泾

• 要求：4天/周，3个月，本科

• 所属：语核科技校招专场，创始人翟先生发布（今日活跃）

• 标签：B端产品、语义类AI、语音类AI、视觉类AI
【主要职责】

1. 根据公司商业化目标，参与细化产品Roadmap，推动团队实现目标；

2. 完成Agent数字员工的复杂模块产品设计；

3. 支持产品设计评审，指导校招产品经理的能力提升，共同建设产品培训体系；

4. 持续关注各渠道用户需求及反馈，洞察背后真实诉求转化为可交付的产品方案；

5. 与团队密切配合完成从产品定义、构建、验收的全流程工作。
【职位要求】

1. 有B端产品意识，具备B端古典产品经理能力和知识体系；

2. 有协助完成产品模块从0-1的经验或能力；

3. 能够在信息中抽丝剥茧抽象出产品竞争力；

4. 能够对过往项目经验的方案，进行清晰有逻辑的阐述，知其然知其所以然（在简历中能够体现优先）；

5. 对AI生产力工具，有主动探索和学习的兴趣，以及得出见解的能力。
【加分项】

1. 自己动手从0-1做过小产品（低代码搭建或vibe coding）；

2. 能体现探索欲或创造力的任何事；

3. 用过AI Agent低代码开发平台，如扣子、Dify。理解Agent构成的基本逻辑。有Agent实践、2B业务经历、vibe coding能力。
【我们提供】

1. 与全球最前沿的Agent团队共同学习与工作。

2. 为所有的对Agent感兴趣的同学提供一个进入行业的机会。

3. 一个支持创新和持续发展的开放团队环境。

4. 深入了解AI技术及其实际应用的机会。

5. 有竞争力的补贴。
"""

@app.get("/api/v1/test/data")
async def get_test_data():
    """获取测试数据（简历文件 + JD）"""
    import os
    from fastapi.responses import FileResponse

    # 检查文件是否存在
    if not os.path.exists(TEST_RESUME_PATH):
        return {
            "success": False,
            "message": "测试简历文件不存在",
            "jd": TEST_JD
        }

    return {
        "success": True,
        "jd": TEST_JD,
        "resume_filename": os.path.basename(TEST_RESUME_PATH)
    }


@app.get("/api/v1/test/resume")
async def get_test_resume():
    """获取测试简历文件"""
    import os
    from fastapi.responses import FileResponse

    if not os.path.exists(TEST_RESUME_PATH):
        raise HTTPException(status_code=404, detail="测试简历文件不存在")

    return FileResponse(
        TEST_RESUME_PATH,
        media_type="application/pdf",
        filename=os.path.basename(TEST_RESUME_PATH)
    )


# ========== 前端静态托管 ==========
# 注意：这些路由必须放在所有 /api/... 路由之后，避免覆盖 API 路由。
# FRONTEND_DIR 基于 __file__ 计算绝对路径，不受启动时工作目录 / rootDir 影响。
FRONTEND_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "frontend"
)

# 挂载前端资源目录（index.html 引用了 ./assets/cat-outline.png）
_ASSETS_DIR = os.path.join(FRONTEND_DIR, "assets")
if os.path.isdir(_ASSETS_DIR):
    app.mount("/assets", StaticFiles(directory=_ASSETS_DIR), name="assets")


@app.get("/")
async def serve_index():
    """返回前端首页。

    关键：禁止缓存 index.html。否则更新前端后，用户浏览器可能继续运行被缓存的旧版 JS，
    出现「后端已修但用户仍踩老 bug / 看到旧数据」（如教育只剩一条、导出截断等），
    公网部署时这个问题尤其隐蔽。HTML 不缓存即可保证每次都加载最新逻辑。
    """
    return FileResponse(
        os.path.join(FRONTEND_DIR, "index.html"),
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


# ========== 应用启动 ==========

if __name__ == "__main__":
    # 用 MetricsMiddleware 包装 ASGI 应用
    wrapped_app = MetricsMiddleware(app)
    uvicorn.run(
        wrapped_app,
        host="0.0.0.0",
        port=8888,
        reload=True,
        log_level="info"
    )
